# root/core/ai_report_service.py

from __future__ import annotations

import os
import re
import json
import math
import hashlib
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import yaml

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass
class DomainMeta:
    domain_key: str                 # domain_0, domain_1...
    acronym: str                    # DG
    name: str                       # Data Governance
    domain_id: str                  # as in flow.yaml domain_id
    sequence: str                   # flow sequence
    dependence: List[str]           # flow dependence list
    qtext: Dict[str, str]           # qid lower -> question text


@dataclass
class DomainScore:
    domain_key: str
    acronym: str
    name: str
    # [(qid, question_text, score_float)]
    question_scores: List[Tuple[str, str, float]]
    avg_raw: float
    avg_floor: int


@dataclass
class DependencyIssue:
    domain_acronym: str
    reference_acronym: str
    dependency_broken: bool
    severity_rationale: str
    scenarios: Dict[str, Any]


@dataclass
class UserMeta:
    user_id: str    
    full_name: str
    email: str


class AIReportService:
    """
    Enterprise report generator aligned with export_service.py logic:
      - Domain/question mapping from FileSystem_Setup.yaml -> flow.yaml + orchestration.yaml (execution_request) + decision_tree.yaml
      - Results scope strictly from 'results' table (dbassessment.xlsx equivalent)

    Scope rules:
      - Non-admin: report for (project_id + user_id)
      - Admin: consolidated report for (project_id + all users), averaging answers per question

    Outputs:
      - DOCX
      - PDF (simple corporate)
    """

    def _likert_label(self, grade: int, lang: str) -> str:
        texts = self._report_texts or {}
        lang = (lang or "us").lower()

        if lang in texts and "likert" in texts[lang]:
            return texts[lang]["likert"].get(int(grade), "")

        if "us" in texts and "likert" in texts["us"]:
            return texts["us"]["likert"].get(int(grade), "")

        return str(grade)
    
    def _load_report_texts(self) -> Dict[str, Dict[str, str]]:
        """
        Loads i18n report texts from:
            root/data/general/report_texts.yaml
        Fallback: empty dict
        """
        path = self.base_dir / "data" / "general" / "report_texts.yaml"

        if not path.exists():
            return {}

        try:
            data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
            return data.get("report_text", {})
        except Exception:
            return {}
            

    def _t(self, key: str, lang: str) -> str:
        lang = (lang or "us").lower()

        texts = self._report_texts or {}

        # idioma específico
        if lang in texts and key in texts[lang]:
            return texts[lang][key]

        # fallback para us
        if "us" in texts and key in texts["us"]:
            return texts["us"][key]

        # fallback final
        return key
        
    
    def __init__(self, base_dir: str, repo: Any):
        self.base_dir = Path(base_dir)
        self.repo = repo        

        # Load report texts from YAML
        self._report_texts = self._load_report_texts()
                
        
    def _force_update_fields_on_open(self, doc: Document):
        settings = doc.settings._element
        update = OxmlElement('w:updateFields')
        update.set(qn('w:val'), 'true')
        settings.append(update)
        
        
    def _override_domain_definition_kv(
        self,
        block: Dict[str, Any],
        domain_acronym: str,
        domain_name: str
    ):
        """
        Se o KV for 'Definição do Domínio',
        substitui o valor por 'DG - Governança de Dados'
        """

        key_norm = (block.get("key") or "").strip().lower()

        if key_norm in ["definição do domínio", "definicao do dominio"]:
            prefix = self._t("domain_definition_prefix", language)
            block["text"] = f"{domain_acronym} - {domain_name}"         

    # =========================================================
    # Public API
    # =========================================================

    def generate_report_docx(
        self,
        project_id: str,
        user_id: str,
        is_admin: bool,
        language: str,
        force_regen: bool = False
    ) -> str:
        out_dir, meta = self._prepare_cache(project_id, user_id, is_admin, language, force_regen)
        docx_path = out_dir / meta["docx_name"]

        if docx_path.exists() and not force_regen:
            force_regen=True
            #return str(docx_path)

        # Build mappings like export_service (filesystem -> flow + orchestration -> decision_tree)
        mapping = self._load_project_mapping(project_id, language=language)
        domain_metas: Dict[str, DomainMeta] = mapping["domain_metas"]
        flow = mapping["flow"]
        orch = mapping["orch"]
        resolved_lang = (language or "us").lower()
        
        T = lambda k: self._t(k, resolved_lang)

        # Extract answers from results table (strict scope)
        answers_by_domain, included_users = self._load_answers(project_id, user_id, is_admin)

        # Build scores strictly for domains present in results/answers and ordered by execution_request/domain_key
        scores = self._compute_scores(domain_metas, answers_by_domain)

        # Dependencies issues knowledge base
        deps_issues = self._load_dependency_issues()

        # Project/User metadata for cover
        project_name = self._get_project_name(project_id)
        user_meta = self._get_user_meta(user_id) if not is_admin else None

        doc = Document()
        self._setup_doc_styles(doc)
        self._force_update_fields_on_open(doc)

        # Cover
        self._add_cover(
            doc=doc,
            project_id=project_id,
            project_name=project_name,
            user_meta=user_meta,
            is_admin=is_admin,
            included_users=included_users,
            language=resolved_lang
        )

        # TOC
        self._add_page_break(doc)
        self._add_heading(doc, T("toc"), level=1)
        self._add_toc_field(doc)
        

        # 1) Results        
        self._add_page_break(doc)
        self._add_heading(doc, T("section_1"), level=1)
        self._add_paragraph(doc, T("section_1_intro_1"))
        self._add_paragraph(doc, T("section_1_intro_2"))
        self._add_results_section(doc, scores, is_admin=is_admin, language=resolved_lang)

        # 2) Dependencies
        #self._add_page_break(doc)
        self._add_heading(doc, T("section_2"), level=1)
        self._add_paragraph(doc, T("section_2_intro"))
        self._add_dependencies_section(doc, domain_metas, scores, deps_issues,resolved_lang)

        # 3) Blueprint
        #self._add_page_break(doc)
        self._add_heading(doc, T("section_3"), level=1)
        self._add_paragraph(doc, T("section_3_intro"))
        
        self._add_blueprint_section(
            doc=doc,
            project_id=project_id,
            scores=scores,
            domain_metas=domain_metas,
            flow=flow,
            language=resolved_lang
        )

        # 4) References
        self._add_page_break(doc)
        self._add_heading(doc, T("section_4"), level=1)
        self._add_references_section(doc, resolved_lang)

        doc.save(str(docx_path))

        # PDF
        self._export_pdf(
            pdf_path=str(out_dir / meta["pdf_name"]),
            project_id=project_id,
            project_name=project_name,
            user_meta=user_meta,
            is_admin=is_admin,
            included_users=included_users,
            language=resolved_lang,
            scores=scores,
            domain_metas=domain_metas,
            deps_issues=deps_issues,
        )

        (out_dir / "report.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

        return str(docx_path)

    def generate_report_pdf(
        self,
        project_id: str,
        user_id: str,
        is_admin: bool,
        language: str,
        force_regen: bool = False
    ) -> str:
        out_dir, meta = self._prepare_cache(project_id, user_id, is_admin, language, force_regen)
        pdf_path = out_dir / meta["pdf_name"]

        if pdf_path.exists() and not force_regen:
            return str(pdf_path)

        _ = self.generate_report_docx(project_id, user_id, is_admin, language, force_regen=force_regen)
        if not pdf_path.exists():
            raise RuntimeError("PDF export failed to generate.")

        return str(pdf_path)

    # =========================================================
    # Cache
    # =========================================================

    def _prepare_cache(self, project_id: str, user_id: str, is_admin: bool, language: str, force_regen: bool):
        language = (language or "us").strip()

        answers_by_domain, included_users = self._load_answers(project_id, user_id, is_admin)

        mapping = self._load_project_mapping(project_id, language=language)
        flow_hash = self._hash_text(yaml.safe_dump(mapping["flow"], sort_keys=True))
        orch_hash = self._hash_text(yaml.safe_dump(mapping["orch"], sort_keys=True))

        fingerprint = self._fingerprint({
            "schema": "report_v2_enterprise",
            "project_id": project_id,
            "user_id": user_id,
            "is_admin": bool(is_admin),
            "language": (language or "us").strip().lower(),
            "answers": answers_by_domain,
            "included_users": [u.user_id for u in included_users],
            "flow_hash": flow_hash,
            "orch_hash": orch_hash,
        })

        reports_root = self.base_dir / "data" / "reports"

        project_name = self._get_project_name(project_id)
        safe_project = re.sub(r'[^A-Za-z0-9_]+', '_', project_name or project_id)
        safe_user = "ALL_USERS" if is_admin else re.sub(r'[^A-Za-z0-9_]+', '_', user_id or "")

        out_dir = reports_root / safe_project / safe_user
        
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")

        # 🔥 CALCULAR VARIÁVEIS ANTES DO DICT
        safe_project = re.sub(r'[^A-Za-z0-9_]+', '_', project_name or "")
        safe_user = "ALL_USERS" if is_admin else re.sub(r'[^A-Za-z0-9_]+', '_', user_id or "")
        seq = datetime.utcnow().strftime("%Y%m%d_%H%M%S")

        meta = {
            "schema": "report_v2_enterprise",
            "project_id": project_id,
            "user_id": user_id,
            "is_admin": bool(is_admin),
            "language": (language or "us").strip().lower(),
            "generated_utc": ts,
            "docx_name": f"{safe_project}_{safe_user}_{seq}.docx",
            "pdf_name": f"{safe_project}_{safe_user}_{seq}.pdf",
            "fingerprint": fingerprint,
        }

        if force_regen:
            for fn in ("report.json", meta["docx_name"], meta["pdf_name"]):
                p = out_dir / fn
                if p.exists():
                    try:
                        p.unlink()
                    except Exception:
                        pass

        return out_dir, meta

    def _hash_text(self, text: str) -> str:
        return hashlib.sha256((text or "").encode("utf-8")).hexdigest()[:16]

    def _fingerprint(self, payload: Any) -> str:
        raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)
        return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]

    # =========================================================
    # Export-service-aligned mapping
    # =========================================================

    def _load_project_mapping(self, project_id: str, language: str) -> Dict[str, Any]:

        fs_path = self._resolve_filesystem_setup(project_id)
        fs_setup = self._safe_load_yaml(fs_path) or {}
        orch_cfg = (fs_setup.get("orchestrator_config") or {})

        flow_rel = orch_cfg.get("main_flow", "flow.yaml")
        orch_rel = orch_cfg.get("main_orchestration", "default_execution.yaml")

        flow_path = self._resolve_project_file(project_id, flow_rel)
        orch_path = self._resolve_project_file(project_id, orch_rel)

        flow = self._safe_load_yaml(flow_path) or {}
        orch = self._safe_load_yaml(orch_path) or {}

        req_list = orch.get("execution_request", []) or []
        domain_flow = flow.get("Domain_flow", []) or []

        # 🔥 idioma estrutural do projeto (para decision_tree)
        project_lang = str(orch.get("language") or "us").strip().lower()
        if not project_lang:
            project_lang = "us"

        domain_metas: Dict[str, DomainMeta] = {}

        for idx, req in enumerate(req_list):
            dom_id = req.get("domain")
            dom_meta = next(
                (d for d in domain_flow if str(d.get("domain_id")) == str(dom_id)),
                None
            ) or {}

            acronym = (dom_meta.get("acronym") or str(dom_id) or f"domain_{idx}").strip()
            name = (dom_meta.get("name") or "").strip()
            sequence = str(dom_meta.get("sequence") or "").strip()
            dependence = dom_meta.get("dependence") or []
            dependence = [str(x) for x in dependence if x is not None]

            files = dom_meta.get("files") or {}
            decision_tree = files.get("decision_tree")

            qtext_map: Dict[str, str] = {}

            if decision_tree:
                tree_path = (
                    self.base_dir
                    / "data"
                    / "domains"
                    / "Language"
                    / project_lang
                    / decision_tree
                )

                tree_data = self._safe_load_yaml(tree_path) or {}
                questions = tree_data.get("questions", {}) or {}

                for qid, qinfo in questions.items():
                    qid_str = str(qid).strip().lower()
                    qtext = (qinfo.get("question") or qinfo.get("text") or "").strip()
                    if qid_str:
                        qtext_map[qid_str] = qtext

            domain_key = f"domain_{idx}"

            domain_metas[domain_key] = DomainMeta(
                domain_key=domain_key,
                acronym=acronym,
                name=name,
                domain_id=str(dom_id) if dom_id is not None else "",
                sequence=sequence,
                dependence=dependence,
                qtext=qtext_map,
            )

        return {
            "fs_path": str(fs_path),
            "flow_path": str(flow_path),
            "orch_path": str(orch_path),
            "flow": flow,
            "orch": orch,
            "project_lang": project_lang,   # 🔥 apenas estrutural
            "domain_metas": domain_metas,
        }

    def _resolve_filesystem_setup(self, project_id: str) -> Path:
        # Prefer project General/FileSystem_Setup.yaml, then project root, then BASE_DIR root
        p1 = self.base_dir / "data" / "projects" / str(project_id) / "General" / "FileSystem_Setup.yaml"
        if p1.exists():
            return p1
        p2 = self.base_dir / "data" / "projects" / str(project_id) / "FileSystem_Setup.yaml"
        if p2.exists():
            return p2
        p3 = self.base_dir / "FileSystem_Setup.yaml"
        if p3.exists():
            return p3
        raise FileNotFoundError("FileSystem_Setup.yaml not found (project or base).")

    def _resolve_project_file(self, project_id: str, rel_path: str) -> Path:
        """
        Resolves a yaml path referenced by filesystem/orchestrator.
        Tries:
          - project General/<rel_path>
          - project root/<rel_path>
          - BASE_DIR/<rel_path>
        """
        rel = str(rel_path or "").strip()
        if not rel:
            raise FileNotFoundError("Empty config path.")

        # If absolute, use directly
        rp = Path(rel)
        if rp.is_absolute() and rp.exists():
            return rp

        candidates = [
            self.base_dir / "data" / "projects" / str(project_id) / "General" / rel,
            self.base_dir / "data" / "projects" / str(project_id) / rel,
            self.base_dir / rel,
        ]
        for c in candidates:
            if c.exists():
                return c

        # last fallback: try lower/upper for "general/General"
        candidates2 = [
            self.base_dir / "data" / "projects" / str(project_id) / "general" / rel,
        ]
        for c in candidates2:
            if c.exists():
                return c

        raise FileNotFoundError(f"Project file not found: {rel}")

    def _safe_load_yaml(self, path: Path) -> Optional[Dict[str, Any]]:
        try:
            return yaml.safe_load(path.read_text(encoding="utf-8"))
        except Exception:
            return None

    def _safe_load_json(self, path: Path) -> Optional[Any]:
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return None

    # =========================================================
    # Results extraction (strict scope)
    # =========================================================

    def _load_answers(
        self,
        project_id: str,
        user_id: str,
        is_admin: bool
    ) -> Tuple[Dict[str, Dict[str, float]], List[UserMeta]]:
        """
        Returns:
          answers_by_domain: {domain_key: {qid: score_float}}
          included_users: users included in report scope

        Scope:
          - Non-admin: only (project_id, user_id)
          - Admin: all users for project_id, averaged per question
        """
        results = self.repo.fetch_all("results") or []
        if not results:
            return {}, []

        # filter rows by scope
        scoped_rows = []
        for r in results:
            if str(r.get("project_id")) != str(project_id):
                continue
            if not is_admin and str(r.get("user_id")) != str(user_id):
                continue
            scoped_rows.append(r)

        if not scoped_rows:
            return {}, []

        # included users metadata
        included_user_ids = sorted({str(r.get("user_id") or "").strip() for r in scoped_rows if str(r.get("user_id") or "").strip()})
        included_users = [self._get_user_meta(uid) for uid in included_user_ids]

        # decrypt and aggregate
        from auth.crypto_service import decrypt_text

        # aggregated[domain_key][qid] -> list of float scores
        aggregated: Dict[str, Dict[str, List[float]]] = {}

        for r in scoped_rows:
            enc = r.get("answers_json_encrypted")
            if not enc:
                continue

            try:
                decrypted = decrypt_text(enc)
                payload = json.loads(decrypted)
            except Exception:
                continue

            answers = payload.get("answers") if isinstance(payload, dict) else None
            if answers is None and isinstance(payload, dict):
                answers = payload

            if not isinstance(answers, dict):
                continue

            for domain_key, qmap in answers.items():
                if not isinstance(qmap, dict):
                    continue

                dk = str(domain_key).strip()
                if not dk:
                    continue

                aggregated.setdefault(dk, {})

                for qid, val in qmap.items():
                    qid_str = str(qid).strip()
                    if not qid_str:
                        continue
                    try:
                        score = float(val)
                    except Exception:
                        continue

                    aggregated[dk].setdefault(qid_str, [])
                    aggregated[dk][qid_str].append(score)

        # finalize
        out: Dict[str, Dict[str, float]] = {}
        for dk, qmap in aggregated.items():
            out[dk] = {}
            for qid, values in qmap.items():
                if not values:
                    continue
                # non-admin will still average (single value)
                out[dk][qid] = sum(values) / float(len(values))

        return out, included_users

    def _get_user_meta(self, user_id: str) -> UserMeta:
        users = self.repo.fetch_all("users") or []
        from auth.crypto_service import decrypt_text

        uid = str(user_id or "").strip()
        for u in users:
            if str(u.get("email_hash") or "").strip() != uid:
                continue
            try:
                full_name = decrypt_text(u.get("full_name_encrypted"))
            except Exception:
                full_name = ""
            try:
                email = decrypt_text(u.get("email_encrypted"))
            except Exception:
                email = ""
            return UserMeta(user_id=uid, full_name=full_name, email=email)

        return UserMeta(user_id=uid, full_name="", email="")

    def _get_project_name(self, project_id: str) -> str:
        projects = self.repo.fetch_all("projects") or []
        pid = str(project_id)
        for p in projects:
            if str(p.get("project_id")) == pid:
                return str(p.get("name") or "").strip()
        return ""

    # =========================================================
    # Scores
    # =========================================================

    def _compute_scores(self, domain_metas: Dict[str, DomainMeta], answers_by_domain: Dict[str, Dict[str, float]]) -> List[DomainScore]:
        """
        Only domains present in answers_by_domain are included.
        Ordering: execution_request order via domain_metas insertion order (domain_0..N created in order).
        """
        scores: List[DomainScore] = []

        for domain_key in domain_metas.keys():
            if domain_key not in answers_by_domain:
                continue

            meta = domain_metas[domain_key]
            qmap = answers_by_domain.get(domain_key, {}) or {}

            qs: List[Tuple[str, str, float]] = []
            for qid, val in qmap.items():
                qid_str = str(qid).strip()
                qtext = meta.qtext.get(qid_str.lower(), "")
                qs.append((qid_str, qtext, float(val)))

            qs.sort(key=lambda x: self._qid_sort_key(x[0]))

            if qs:
                vals = [v for _, _, v in qs]
                avg_raw = sum(vals) / float(len(vals))
                avg_floor = int(math.floor(avg_raw))
            else:
                avg_raw = 0.0
                avg_floor = 0

            scores.append(DomainScore(
                domain_key=domain_key,
                acronym=meta.acronym,
                name=meta.name,
                question_scores=qs,
                avg_raw=avg_raw,
                avg_floor=avg_floor
            ))

        return scores

    def _qid_sort_key(self, qid: str):
        s = str(qid).strip().lower()
        m = re.search(r"(\d+)", s)
        return int(m.group(1)) if m else 9999

    # =========================================================
    # Dependencies
    # =========================================================

    def _load_dependency_issues(self) -> List[DependencyIssue]:
        p = self.base_dir / "data" / "domains" / "us" / "Dependencies_inconsistencies_theory_cluster_output.json"
        data = self._safe_load_json(p)
        if not data:
            return []

        if isinstance(data, list):
            items = data
        elif isinstance(data, dict):
            items = data.get("items") or data.get("data") or []
        else:
            items = []

        out: List[DependencyIssue] = []
        for it in items:
            if not isinstance(it, dict):
                continue
            out.append(DependencyIssue(
                domain_acronym=str(it.get("domain_acronym") or "").strip(),
                reference_acronym=str(it.get("reference_acronym") or "").strip(),
                dependency_broken=bool(it.get("dependency_broken", False)),
                severity_rationale=str(it.get("severity_rationale") or "").strip(),
                scenarios=it.get("scenarios") or {}
            ))
        return out

    # =========================================================
    # Blueprint: Action + DEMO
    # =========================================================

    def _triggered_action_code(self, acronym: str, maturity_grade: int) -> str:
        return f"{acronym}-{maturity_grade + 1:02d}"

    def _load_demo_item_for_action(self, acronym: str, action_code: str, language: str):

        lang = (language or "us").lower()

        p = (
            self.base_dir
            / "data"
            / "domains"
            / lang
            / f"{acronym}_theory_demo_output_PATCHED.json"
        )

        if not p.exists():
            return None

        data = self._safe_load_json(p)
        if not data:
            return None

        items = data.get("items") if isinstance(data, dict) else data
        if not isinstance(items, list):
            return None

        for it in items:
            if str(it.get("action_code") or "").strip() == action_code:
                return it

        return None

    # =========================================================
    # DOCX helpers
    # =========================================================

    def _render_text_with_bold_prefix_if_colon(
        self,
        doc: Document,
        text: str
    ):
        """
        Se a linha tiver formato 'Tema: Conteúdo',
        coloca o texto antes dos dois pontos em negrito.
        """

        if ":" not in text:
            self._add_paragraph(doc, text)
            return

        prefix, suffix = text.split(":", 1)

        # Evitar quebrar URLs ou textos estranhos
        if len(prefix.strip()) < 2:
            self._add_paragraph(doc, text)
            return

        p = doc.add_paragraph()
        p.add_run(prefix.strip() + ": ").bold = True
        p.add_run(suffix.strip())
    
    def _setup_doc_styles(self, doc: Document):

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        sizes = {
            1: 20,
            2: 16,
            3: 14,
            4: 12
        }

        for lvl, size in sizes.items():
            st = doc.styles[f"Heading {lvl}"]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.bold = True

    def _add_page_break(self, doc: Document):
        doc.add_page_break()

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        doc.add_heading(text, level=level)

    def _add_paragraph(self, doc: Document, text: str):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        
    def _resolve_action_catalog_path(self, domain_acronym: str, language: str) -> Optional[Path]:
        """
        Tries to locate <DOMAIN>_action_catalog.yaml in common repo locations.
        We keep this very defensive to avoid breaking existing behavior.
        """
        dom = (domain_acronym or "").strip()
        lang = (language or "us").strip().lower()

        if not dom:
            return None

        filenames = [
            f"{dom}_action_catalog.yaml",
            f"{dom}_action_catalog.yml",
        ]

        candidates: List[Path] = []

        # Most common in your project structure (same pattern used for decision_tree)
        for fn in filenames:
            candidates.append(self.base_dir / "data" / "domains" / "Language" / lang / fn)
            candidates.append(self.base_dir / "data" / "domains" / "Language" / "us" / fn)

            # Legacy / alternate layouts
            candidates.append(self.base_dir / "data" / "domains" / lang / fn)
            candidates.append(self.base_dir / "data" / "domains" / "us" / fn)

            # Last resort (project root)
            candidates.append(self.base_dir / fn)

        for p in candidates:
            if p.exists():
                return p

        return None

    def _load_action_catalog(self, domain_acronym: str, language: str) -> Optional[Dict[str, Any]]:
        p = self._resolve_action_catalog_path(domain_acronym, language)
        if not p:
            return None
        data = self._safe_load_yaml(p)
        return data if isinstance(data, dict) else None

    def _add_procedure_elements_of_example(
        self,
        doc: Document,
        domain_acronym: str,
        action_code: str,
        language: str
    ):
        """
        Adds a new subsection "Elementos de exemplo" listing ALL procedures
        from action_catalog.yaml for the given domain+action.
        Format: titles before ':' in bold.
        """
        catalog = self._load_action_catalog(domain_acronym, language)
        if not catalog:
            return

        action = (catalog.get("action_catalog") or {}).get(action_code) or {}
        procedures = action.get("procedures") or []
        if not isinstance(procedures, list) or not procedures:
            return

        # New subsection
        self._add_heading(doc, self._t("example_elements_title", language), level=4)

        def add_kv(key: str, val: str):
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(val or "")

        for i, proc in enumerate(procedures, start=1):
            if not isinstance(proc, dict):
                continue

            # Small separation between procedures (not a page break)
            if i > 1:
                doc.add_paragraph("")

            add_kv(self._t("procedure_number", language), str(proc.get("number") or "").strip())
            add_kv(self._t("procedure_name", language), str(proc.get("name") or "").strip())
            add_kv(self._t("procedure_prerequisite", language), str(proc.get("prerequisite") or "").strip())
            add_kv(
                self._t("procedure_deliverable", language),
                str(proc.get("deliverable") or "").strip()
            )

            recs = proc.get("recommendations") or []
            if isinstance(recs, list) and recs:
                p = doc.add_paragraph()
                p.add_run(self._t("procedure_recommendations", language)).bold = True
                for r in recs:
                    rr = str(r or "").strip()
                    if rr:
                        doc.add_paragraph(rr, style="List Bullet")

            note_value = proc.get("note") or proc.get("notes")
            if note_value:
                p = doc.add_paragraph()
                p.add_run(self._t("procedure_notes", language)).bold = True
                if isinstance(note_value, list):
                    for n in note_value:
                        nn = str(n or "").strip()
                        if nn:
                            doc.add_paragraph(nn, style="List Bullet")
                else:
                    txt = str(note_value).strip()
                    if txt:
                        doc.add_paragraph(txt, style="List Bullet")
    
    
    def _add_toc_field(self, doc: Document):
        p = doc.add_paragraph()
        r = p.add_run()

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_separate = OxmlElement('w:fldChar')
        fld_separate.set(qn('w:fldCharType'), 'separate')

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        r._r.append(fld_begin)
        r._r.append(instr)
        r._r.append(fld_separate)
        r._r.append(fld_end)

        doc.add_paragraph("")

        #doc.add_paragraph("Update the Table of Contents in Word (Right click → Update field) after opening the document.")

    def _add_cover(
        self,
        doc: Document,
        project_id: str,
        project_name: str,
        user_meta: Optional[UserMeta],
        is_admin: bool,
        included_users: List[UserMeta],
        language: str
    ):
        
        T = lambda k: self._t(k, language)
        
        title = doc.add_paragraph()
        run = title.add_run(T("cover_title_admin") if is_admin else T("cover_title_user"))
        run.bold = True
        run.font.size = Pt(28)

        subtitle = doc.add_paragraph()
        run2 = subtitle.add_run(T("cover_subtitle"))
        run2.font.size = Pt(16)

        doc.add_paragraph("")

        p1 = doc.add_paragraph()
        p1.add_run(f"{T('project')} ").bold = True
        p1.add_run(project_name or "")
        if project_name:
            p1.add_run(" ")
        p1.add_run(f"(ID: {project_id})")

        if is_admin:
            p2 = doc.add_paragraph()            
            p2.add_run(f"{T('scope')} ").bold = True
            p2.add_run(T("scope_all_users"))

            p3 = doc.add_paragraph()            
            p3.add_run(f"{T('users_included')} ").bold = True
            p3.add_run(str(len(included_users)))

        else:
            p2 = doc.add_paragraph()            
            p2.add_run(f"{T('assessed_user')} ").bold = True
            
            if user_meta and (user_meta.full_name or user_meta.email):
                p2.add_run(user_meta.full_name or "")
                if user_meta.email:
                    p2.add_run(f" ({user_meta.email})")
                p2.add_run(f" [ID: {user_meta.user_id}]")
            else:
                p2.add_run(f"[ID: {user_meta.user_id if user_meta else ''}]")

        p4 = doc.add_paragraph()        
        p4.add_run(f"{T('report_language')} ").bold = True
        p4.add_run(str(language))

        p5 = doc.add_paragraph()        
        p5.add_run(f"{T('generated')} ").bold = True
        p5.add_run(datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"))

    # =========================================================
    # DOCX sections
    # =========================================================

    def _add_results_section(self, doc: Document, scores: List[DomainScore], is_admin: bool, language: str):
        
        T = lambda k: self._t(k, language)
        
        doc.add_paragraph("")        
        self._add_heading(doc, T("domain_summary_title"), level=2)

        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = T("table_domain")
        hdr[1].text = T("table_domain_name")
        hdr[2].text = T("table_avg")
        hdr[3].text = T("table_final_grade")
        hdr[4].text = T("table_maturity_level")

        for s in scores:
            row = table.add_row().cells
            row[0].text = s.acronym
            row[1].text = s.name
            row[2].text = f"{s.avg_raw:.2f}"
            row[3].text = str(s.avg_floor)
            row[4].text = self._likert_label(s.avg_floor, language)

        for s in scores:
            doc.add_paragraph("")
            self._add_heading(doc, f"1.2 {s.acronym} · {s.name}", level=2)

            self._add_paragraph(
                doc,
                f"{T('domain_maturity_rule')} {s.avg_floor} ({self._likert_label(s.avg_floor,language)})."
            )

            if not s.question_scores:                
                self._add_paragraph(doc, T("no_scored_questions"))
                continue

            t = doc.add_table(rows=1, cols=4)
            h = t.rows[0].cells
            h[0].text = T("table_question")
            h[1].text = T("table_score")
            h[2].text = T("table_nearest_level")
            h[3].text = T("table_notes")

            for qid, qtext, val in s.question_scores:
                r = t.add_row().cells
                r[0].text = qtext or qid
                r[1].text = f"{val:.2f}" if is_admin else str(int(round(val)))
                nearest = int(round(val))
                if nearest < 0:
                    nearest = 0
                if nearest > 5:
                    nearest = 5
                r[2].text = self._likert_label(nearest, language)
                r[3].text = ""
            
            self._add_paragraph(doc, T("scoring_rule"))

    def _add_dependencies_section(
        self,
        doc: Document,
        domain_metas: Dict[str, DomainMeta],
        scores: List[DomainScore],
        issues: List[DependencyIssue],
        language: str
    ):
            
        T = lambda k: self._t(k, language)
        
        # Only evaluate dependencies for domains in scope
        in_scope_acr = {s.acronym for s in scores}
        score_by_acr = {s.acronym: s.avg_floor for s in scores}

        # Build acronym -> name lookup from domain_metas
        acr_to_name = {m.acronym: m.name for m in domain_metas.values()}
        
        self._add_heading(doc, T("declared_dependencies_heading"), level=2)

        for m in domain_metas.values():
            if m.acronym not in in_scope_acr:
                continue

            self._add_heading(doc, f"{m.acronym} · {m.name}", level=3)

            dep = m.dependence or []
            if not dep:                
                self._add_paragraph(doc, T("no_dependencies"))
                continue

            # Your rule: last element is the domain itself; previous are dependencies
            deps_only = []

            if len(dep) >= 2:
                deps_only = dep[:-1]
            else:
                deps_only = dep

            # remover autorreferência
            deps_only = [
                x for x in deps_only
                if not any(
                    str(mm.domain_id) == str(x) and mm.acronym == m.acronym
                    for mm in domain_metas.values()
                )
            ]

            if not deps_only:                
                self._add_paragraph(doc, T("no_dependencies"))
                continue
            
            self._add_paragraph(doc, T("declared_dependencies_label"))
            
            for x in deps_only:
                # dependence numbers represent domain_id. We need resolve to acronym via domain_metas domain_id
                dep_acr = None
                dep_name = ""
                for mm in domain_metas.values():
                    if str(mm.domain_id) == str(x):
                        dep_acr = mm.acronym
                        dep_name = mm.name
                        break

                if not dep_acr:                    
                    doc.add_paragraph(f"- {T('unknown_domain')} (domain_id={x})", style="List Bullet")
                    continue

                dep_grade = score_by_acr.get(dep_acr)
                if dep_grade is None:                    
                    doc.add_paragraph(f"- {dep_acr} · {dep_name} ({T('not_evaluated_scope')})", style="List Bullet")
                else:
                    doc.add_paragraph(f"- {dep_acr} · {dep_name} (grade {dep_grade} · {self._likert_label(dep_grade,language)})", style="List Bullet")

        #self._add_page_break(doc)        
        self._add_heading(doc, T("detected_breaks"), level=2)

        relevant = [x for x in issues if x.domain_acronym in in_scope_acr]
        broken = [x for x in relevant if x.dependency_broken]

        if not relevant:
            self._add_paragraph(doc, T("no_dependency_records"))
            return

        if not broken:
            self._add_paragraph(doc, T("no_dependency_breaks"))
            return

        for it in broken:
            dname = acr_to_name.get(it.domain_acronym, "")
            rname = acr_to_name.get(it.reference_acronym, "")
            
            title = f"{it.domain_acronym} · {dname} {T('depends_on')} {it.reference_acronym} · {rname}"
            
            self._add_heading(doc, title, level=3)

            if it.severity_rationale:
                self._add_paragraph(doc, f"{T('severity_rationale_label')} {it.severity_rationale}")

            if isinstance(it.scenarios, dict) and it.scenarios:
                for sk in sorted(it.scenarios.keys()):
                    sc = it.scenarios.get(sk) or {}
                    if not isinstance(sc, dict):
                        continue

                    sc_title = sc.get("comparison") or sk                    
                    self._add_heading(doc, f"{T('scenario_label')} {sc_title}", level=3)

                    for field in ("whynot_text", "whatcauses_text", "howtofix_text", "analysis_text"):
                        txt = sc.get(field)
                        if txt and isinstance(txt, str):
                            self._add_paragraph(doc, txt.strip())

    
    def _norm_title(self, s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip().lower())
        
    def _clean_placeholders(self, text: str) -> str:
        """
        Remove tokens internos tipo __ACR__ACR12____
        sem afetar conteúdo legítimo.
        """
        if not text:
            return text

        # remove padrões do tipo __XXXX__XXXX____
        text = re.sub(r"__[^_]+__[^_]+____", "", text)

        return text.strip()

    def _is_domain_context_section(self, sec_title_norm: str) -> bool:
        # "Domain and Context Framing" / "Domínio e Contexto"
        return ("domain and context" in sec_title_norm) or ("domínio e contexto" in sec_title_norm) or ("dominio e contexto" in sec_title_norm)

    def _is_action_definition_section(self, sec_title_norm: str) -> bool:
        # "Action Definition" / "Definição de Ação"
        return ("action definition" in sec_title_norm) or ("definição de ação" in sec_title_norm) or ("definicao de acao" in sec_title_norm)

    def _is_procedure_definition_section(self, sec_title_norm: str) -> bool:
        # "Procedure Definition" / "Definição de Procedimento"
        return ("procedure definition" in sec_title_norm) or ("definição de procedimento" in sec_title_norm) or ("definicao de procedimento" in sec_title_norm)

    def _is_step_by_step_subsection(self, sub_title_norm: str) -> bool:
        # "Step-by-Step Execution Structure" / "Estrutura de Execução Passo a Passo"
        return ("step-by-step" in sub_title_norm) or ("passo a passo" in sub_title_norm) or ("execução passo" in sub_title_norm) or ("execucao passo" in sub_title_norm)

    def _load_domain_tree_and_catalog(
        self,
        project_id: str,
        domain_acronym: str,
        flow: Dict[str, Any],
        language: str
    ) -> Tuple[Optional[Dict[str, Any]], Optional[Dict[str, Any]]]:
        """
        Resolve paths no padrão do projeto (data/projects/<id>/Domains/<lang>/...)
        Usando flow.yaml Domain_flow -> files: decision_tree / action_catalog.
        """
        lang = (language or "us").strip().lower() or "us"

        domain_flow = flow.get("Domain_flow", []) or []
        dom = next((d for d in domain_flow if str(d.get("acronym") or "").strip().upper() == str(domain_acronym).strip().upper()), None) or {}
        files = dom.get("files") or {}

        decision_tree_rel = files.get("decision_tree")
        action_catalog_rel = files.get("action_catalog")

        base = self.base_dir / "data" / "projects" / str(project_id) / "Domains"

        candidates_lang = [lang, "us"]
        tree_data = None
        catalog_data = None

        for lg in candidates_lang:
            if decision_tree_rel and tree_data is None:
                p = base / lg / str(decision_tree_rel)
                if p.exists():
                    tree_data = self._safe_load_yaml(p) or None

            if action_catalog_rel and catalog_data is None:
                p = base / lg / str(action_catalog_rel)
                if p.exists():
                    catalog_data = self._safe_load_yaml(p) or None

            if tree_data is not None and catalog_data is not None:
                break

        return tree_data, catalog_data

    def _first_question_node(self, tree_data: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        if not isinstance(tree_data, dict):
            return None
        questions = tree_data.get("questions") or {}
        if not isinstance(questions, dict) or not questions:
            return None

        # pega a primeira por ordem Q1, Q2... se possível
        def qkey(k: str) -> int:
            m = re.search(r"(\d+)", str(k))
            return int(m.group(1)) if m else 9999

        first_key = sorted(questions.keys(), key=qkey)[0]
        qnode = questions.get(first_key)
        return qnode if isinstance(qnode, dict) else None

    def _domain_description(self, tree_data: Optional[Dict[str, Any]]) -> str:
        if not isinstance(tree_data, dict):
            return ""
        dom = tree_data.get("domain") or {}
        if isinstance(dom, dict):
            return str(dom.get("description") or "").strip()
        return ""

    def _domain_objectives(self, tree_data: Optional[Dict[str, Any]]) -> List[str]:
        if not isinstance(tree_data, dict):
            return []
        questions = tree_data.get("questions") or {}
        if not isinstance(questions, dict):
            return []
        out = []
        for _, qinfo in questions.items():
            if not isinstance(qinfo, dict):
                continue
            obj = str(qinfo.get("objective") or "").strip()
            if obj:
                out.append(obj)
        return out

    def _action_from_tree_for_score(self, tree_data: Optional[Dict[str, Any]], score_floor: int) -> Tuple[str, str]:
        """
        Retorna (action_code, description) a partir do score_action_mapping do decision_tree.
        Assume consistência do mapping entre questions.
        """
        qnode = self._first_question_node(tree_data)
        if not isinstance(qnode, dict):
            return "", ""

        mapping = qnode.get("score_action_mapping") or {}
        if not isinstance(mapping, dict):
            return "", ""

        action_node = mapping.get(score_floor)
        if not isinstance(action_node, dict):
            # às vezes o YAML serializa keys como string
            action_node = mapping.get(str(score_floor))

        if not isinstance(action_node, dict):
            return "", ""

        ac = str(action_node.get("action_code") or "").strip()
        desc = str(action_node.get("description") or "").strip()
        return ac, desc

    def _find_action_in_catalog(self, catalog_data: Optional[Dict[str, Any]], action_code: str) -> Optional[Dict[str, Any]]:
        if not isinstance(catalog_data, dict):
            return None
        actions = catalog_data.get("actions")
        if isinstance(actions, dict):
            node = actions.get(action_code)
            return node if isinstance(node, dict) else None
        if isinstance(actions, list):
            for it in actions:
                if isinstance(it, dict) and str(it.get("action_code") or "").strip() == action_code:
                    return it
        return None

    def _add_maturity_levels_section(
        self,
        doc: Document,
        tree_data: Dict[str, Any],
        language: str
    ):
        """
        Renderiza maturity_scale completo do decision_tree.yaml
        """
        
        T = lambda k: self._t(k, language)
        
        maturity_scale = tree_data.get("maturity_scale")
        if not isinstance(maturity_scale, dict):
            return

        def sort_key(k):
            try:
                return int(k)
            except Exception:
                return 999

        for level_key in sorted(maturity_scale.keys(), key=sort_key):

            level_data = maturity_scale.get(level_key)
            if not isinstance(level_data, dict):
                continue

            maturity_level = level_data.get("maturity_level", "")
            meaning = level_data.get("meaning", "")
            interpretation = level_data.get("interpretation_for_assessment", "")

            # Heading por nível
            self._add_heading(
                doc,
                f"{T('maturity_level_prefix')} {level_key} – {maturity_level}",
                level=3
            )

            if meaning:
                self._add_paragraph(doc, meaning)

            if interpretation:
                self._add_paragraph(doc, interpretation)

    def _inject_domain_and_context(self, doc: Document, tree_data: Optional[Dict[str, Any]]):
        """
        3.x.3 Definição de Domínio e Contexto:
          - 1 parágrafo: decision_tree.domain.description
          - bullets: todos os objectives das questions
        """
        desc = self._domain_description(tree_data)
        if desc:
            self._add_paragraph(doc, desc)

        objectives = self._domain_objectives(tree_data)
        if objectives:
            for obj in objectives:
                doc.add_paragraph(obj, style="List Bullet")

    def _inject_action_definition(self, doc: Document, tree_data: Optional[Dict[str, Any]], score_floor: int):
        """
        3.x.4 Definição de Ação:
          - 1 parágrafo: action_code + description vindo do score_action_mapping do decision_tree
        """
        ac, desc = self._action_from_tree_for_score(tree_data, score_floor)
        if not ac and not desc:
            return

        if desc:
            self._add_paragraph(doc, f"{ac}: {desc}" if ac else desc)
        else:
            self._add_paragraph(doc, f"{ac}")
    
                
    def _add_maturity_levels_block(
        self,
        doc: Document,
        project_id: str,
        scores: List[DomainScore],
        flow: Dict[str, Any],
        language: str
    ):
        """
        Renderiza 3.1 Maturity Levels com base no maturity_scale
        do decision_tree.yaml do primeiro domínio.
        Não altera nada do fluxo existente.
        """

        if not scores:
            return

        domain_acronym = scores[0].acronym

        tree_data, _ = self._load_domain_tree_and_catalog(
            project_id=project_id,
            domain_acronym=domain_acronym,
            flow=flow,
            language=language
        )

        if not isinstance(tree_data, dict):
            return

        maturity_scale = tree_data.get("maturity_scale")
        if not isinstance(maturity_scale, dict):
            return

        # Heading principal
        self._add_heading(doc, T("maturity_levels_title"), level=2)

        # Ordenar níveis numericamente
        def sort_key(k):
            try:
                return int(k)
            except Exception:
                return 999

        for level_key in sorted(maturity_scale.keys(), key=sort_key):

            level_data = maturity_scale.get(level_key)
            if not isinstance(level_data, dict):
                continue

            maturity_level = level_data.get("maturity_level", "")
            meaning = level_data.get("meaning", "")
            interpretation = level_data.get("interpretation_for_assessment", "")

            # Subheading por nível
            self._add_heading(
                doc,
                f"{T('maturity_level_prefix')} {level_key} – {maturity_level}",
                level=3
            )

            if meaning:
                self._add_paragraph(doc, meaning)

            if interpretation:
                self._add_paragraph(doc, interpretation)
                
    
    def _add_blueprint_section(
        self,
        doc: Document,
        project_id: str,
        scores: List[DomainScore],
        domain_metas: Dict[str, DomainMeta],
        flow: Dict[str, Any],
        language: str
    ):
        T = lambda k: self._t(k, language)

        self._add_paragraph(doc, T("blueprint_context"))

        # =========================
        # 3.1 Maturity Levels
        # =========================
        maturity_tree = None

        if scores:
            maturity_tree, _ = self._load_domain_tree_and_catalog(
                project_id=project_id,
                domain_acronym=scores[0].acronym,
                flow=flow,
                language=language
            )

        if isinstance(maturity_tree, dict) and maturity_tree.get("maturity_scale"):
            self._add_heading(doc, T("maturity_levels_title"), level=2)
            self._add_maturity_levels_section(doc, maturity_tree, language)
            
        # =========================
        # 3.2+ Domains
        # =========================
        for idx, s in enumerate(scores, start=2):

            chapter_number = f"3.{idx}"

            self._add_heading(
                doc,
                f"{chapter_number} {s.acronym} · {s.name}",
                level=2
            )

            self._add_paragraph(
                doc,
                f"{T('current_maturity')} {s.avg_floor} ({self._likert_label(s.avg_floor, language)})."
            )

            action_code = self._triggered_action_code(s.acronym, s.avg_floor)

            item = self._load_demo_item_for_action(
                s.acronym,
                action_code,
                language
            )

            title = ""
            demo = ""

            if item:
                title = str(
                    item.get("action_title")
                    or item.get("title")
                    or item.get("name")
                    or ""
                ).strip()

                demo = str(item.get("demo") or "").strip()

            if title:
                self._add_paragraph(doc, f"{T('triggered_action')} {action_code} · {title}")
            else:
                self._add_paragraph(doc, f"{T('triggered_action_code')} {action_code}")

            if not demo:
                self._add_paragraph(doc, T("no_demo"))
                continue

            # Subtítulo estruturado (aparece no sumário)
            self._add_heading(
                doc,
                f"{chapter_number} {T('procedure_pack_full')}",
                level=3
            )

            sections = self._parse_demo_sections(demo)

            if not sections:
                self._add_paragraph(doc, demo)
                continue

            render_example_block = True
            
            # carregar decision_tree e action_catalog do projeto (path estilo renderer_assessment)
            tree_data, catalog_data = self._load_domain_tree_and_catalog(
                project_id=project_id,
                domain_acronym=s.acronym,
                flow=flow,
                language=language
            )

            # localizar action do catálogo
            catalog_action = self._find_action_in_catalog(catalog_data, action_code)            
            
            
            # loop do DEMO
            for sec_index, sec in enumerate(sections, start=1):

                section_number = f"{chapter_number}.{sec_index}"

                self._add_heading(
                    doc,
                    f"{section_number} {sec['title']}",
                    level=3
                )

                sec_title_norm = self._norm_title(sec.get("title", ""))

                # flags para injeções pontuais
                inject_domain_context = self._is_domain_context_section(sec_title_norm)
                inject_action_definition = self._is_action_definition_section(sec_title_norm)
                inject_procedure_definition = self._is_procedure_definition_section(sec_title_norm)

                injected_procedures_after_step_by_step = False

                for block in sec["blocks"]:

                    if block["kind"] == "subsection":
                        self._add_heading(doc, block["title"], level=4)

                        for item in block["items"]:
                            if item["type"] == "bullet":
                                doc.add_paragraph(item["text"], style="List Bullet")
                            elif item["type"] == "kv":
                                p = doc.add_paragraph()
                                p.add_run(item["key"] + ": ").bold = True
                                p.add_run(item["text"])
                            else:
                                self._render_text_with_bold_prefix_if_colon(
                                    doc,
                                    item["text"]
                                )

                        # -------------------------------------------------
                        # 4a alteração: após "Estrutura de Execução Passo a Passo"
                        # adicionar novo subtopico "Elementos de exemplo"
                        # apenas uma vez por domínio (Q1 como referência)
                        # -------------------------------------------------

                        sec_title = str(sec.get("title") or "").strip()
                        is_procedure_definition = (
                            "Definição" in sec_title and "Procedimento" in sec_title
                        )

                        if (
                            render_example_block
                            and is_procedure_definition
                            and str(block.get("title") or "").strip() == "Estrutura de Execução Passo a Passo"
                        ):
                            self._add_procedure_elements_of_example(
                                doc=doc,
                                domain_acronym=s.acronym,
                                action_code=action_code,
                                language=language
                            )
                            render_example_block = False

                    elif block["kind"] == "kv":

                        # --- override apenas para Definição do Domínio ---
                        self._override_domain_definition_kv(
                            block=block,
                            domain_acronym=s.acronym,
                            domain_name=s.name
                        )
                        # --- fim override ---

                        p = doc.add_paragraph()
                        p.add_run(block["key"] + ": ").bold = True
                        p.add_run(block["text"])

                    elif block["kind"] == "bullet":
                        doc.add_paragraph(block["text"], style="List Bullet")

                    else:
                        self._render_text_with_bold_prefix_if_colon(
                            doc,
                            block["text"]
                        )

                # injeções no final da seção (3.x.3 e 3.x.4)
                if inject_domain_context:
                    self._inject_domain_and_context(doc, tree_data)

                if inject_action_definition:
                    self._inject_action_definition(doc, tree_data, s.avg_floor)

    def _parse_demo_sections(self, demo: str) -> List[Dict[str, Any]]:
        text = (demo or "").replace("\r\n", "\n").strip()
        if not text:
            return []

        # Detecta títulos no formato "1. Title"
        pat = re.compile(r"(?m)^\s*(\d+)\.\s+(.+?)\s*$")
        matches = list(pat.finditer(text))
        if not matches:
            return []

        chunks = []
        for i, m in enumerate(matches):
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            chunks.append(text[start:end].strip())

        out = []
        for ch in chunks:
            lines = ch.split("\n")

            raw_title = lines[0].strip()

            # 🔥 REMOVE prefixo numérico tipo "1. ", "2. ", etc.
            title = re.sub(r"^\s*\d+\.\s*", "", raw_title).strip()

            body = "\n".join(lines[1:]).strip()
            blocks = self._parse_demo_body_blocks(body)

            out.append({
                "title": title,
                "blocks": blocks
            })

        return out

    def _parse_demo_body_blocks(self, body: str) -> List[Dict[str, Any]]:
        if not body:
            return []

        blocks = []
        current_subsection = None

        for raw in body.split("\n"):
            if not raw.strip():
                continue

            line = re.sub(r"^\s*\d+(\.\d+)*\.?\s*", "", raw.strip())
            line = self._clean_placeholders(line)

            # 🔹 Subsection header (termina com :)
            if line.endswith(":") and not ":" in line[:-1]:
                current_subsection = line[:-1].strip()
                blocks.append({
                    "kind": "subsection",
                    "title": current_subsection,
                    "items": []
                })
                continue

            # 🔹 Bullet explícito
            if line.startswith("- "):
                text = line[2:].strip()

                if blocks and blocks[-1]["kind"] == "subsection":
                    blocks[-1]["items"].append({"type": "bullet", "text": text})
                else:
                    blocks.append({"kind": "bullet", "text": text})
                continue

            # 🔹 KV
            m = re.match(r"^([A-Za-z][A-Za-z\s]+):\s*(.+)$", line)
            if m:
                key = m.group(1).strip()
                val = m.group(2).strip()

                if blocks and blocks[-1]["kind"] == "subsection":
                    blocks[-1]["items"].append({"type": "kv", "key": key, "text": val})
                else:
                    blocks.append({"kind": "kv", "key": key, "text": val})
                continue

            # 🔹 Texto simples
            if blocks and blocks[-1]["kind"] == "subsection":
                blocks[-1]["items"].append({"type": "text", "text": line})
            else:
                blocks.append({"kind": "text", "text": line})

        return blocks

    def _add_references_section(self, doc: Document, language: str):
        T = lambda k: self._t(k, language)
        refs = [
            "DAMA International. (n.d.). DAMA-DMBOK: Data Management Body of Knowledge (2nd ed.).",
            "EDM Council. (n.d.). DCAM: Data Management Capability Assessment Model.",
            "CMMI Institute. (n.d.). CMMI Model.",
            "European Union. (2016). General Data Protection Regulation (GDPR) (EU) 2016/679.",
            "DOMMx publication (Zenodo record: 18020434).",
            "SLR DOMMx publication (IEEE Xplore document: 11223201).",
        ]
        for r in refs:
            doc.add_paragraph(r, style="List Bullet")

    # =========================================================
    # PDF export
    # =========================================================

    def _export_pdf(
        self,
        pdf_path: str,
        project_id: str,
        project_name: str,
        user_meta: Optional[UserMeta],
        is_admin: bool,
        included_users: List[UserMeta],
        language: str,
        scores: List[DomainScore],
        domain_metas: Dict[str, DomainMeta],
        deps_issues: List[DependencyIssue],
    ):
        # PDF temporarily disabled
        return

    def _pdf_draw_wrapped(self, c, text: str, x: float, y: float, max_w: float, leading: float) -> float:
        words = (text or "").split()
        if not words:
            return y

        line = ""
        for w in words:
            trial = (line + " " + w).strip()
            if c.stringWidth(trial, c._fontname, c._fontsize) <= max_w:
                line = trial
            else:
                c.drawString(x, y, line)
                y -= leading
                line = w

        if line:
            c.drawString(x, y, line)
            y -= leading

        return y