import json
import time

from pathlib import Path
from typing import Dict, Any, List, Optional

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None


REPORT_ENGINE_VERSION = "DEMO_STATIC_V4_FULL_COMPAT"


class AIReportService:

    def __init__(
        self,
        repo=None,
        base_dir: str = ".",
        catalog_lang: str = "us",
        **kwargs
    ):
        self.repo = repo
        self.base_dir = Path(base_dir)
        self.catalog_lang = catalog_lang


    # ==========================================================
    # PUBLIC METHOD EXPECTED BY APP
    # ==========================================================

    def generate_report_docx(
        self,
        project_id: str,
        user_id: str = None,
        is_admin: bool = False,
        language: str = "pt",
        force_regen: bool = False,
        **kwargs
    ) -> str:
            """
            Compatível com auth_service.py atual.
            """

            # 1️⃣ Buscar resultados do projeto
            results = self.repo.fetch_all("results") or []

            project_results = [
                r for r in results
                if str(r.get("project_id")) == str(project_id)
            ]

            if not project_results:
                raise ValueError("No results found for this project.")

            # 2️⃣ Montar report_json mínimo estruturado
            domains = []

            for r in project_results:
                domains.append({
                    "acronym": r.get("domain_acronym"),
                    "name": r.get("domain_name"),
                    "level": r.get("level"),
                    "level_entries": r.get("level_entries", [])
                })

            report_json = {
                "labels": {
                    "report_title": "DOMMx Final Assessment Report"
                },
                "meta": {
                    "scope": f"Project {project_id}",
                    "generated_at": time.strftime("%Y-%m-%d %H:%M:%S")
                },
                "domains": domains
            }

            # 3️⃣ Definir output
            output_dir = self.base_dir / "output"
            output_dir.mkdir(exist_ok=True)

            output_path = output_dir / f"DOMMx_Report_{project_id}.docx"

            # 4️⃣ Gerar docx
            self._generate_docx(report_json, str(output_path))

            return str(output_path)

    # ==========================================================
    # INTERNAL DOCX GENERATOR
    # ==========================================================

    def _generate_docx(
        self,
        report_json: Dict[str, Any],
        output_path: str
    ) -> None:

        if Document is None:
            raise RuntimeError("python-docx not installed")

        doc = Document()

        self._render_cover(doc, report_json)
        self._render_toc(doc)
        self._render_results_section(doc, report_json)
        self._render_demo_section(doc, report_json)

        doc.save(output_path)

    # ==========================================================
    # COVER
    # ==========================================================

    def _render_cover(self, doc: Document, report_json: Dict[str, Any]):

        labels = report_json.get("labels", {}) or {}
        meta = report_json.get("meta", {}) or {}

        title = labels.get("report_title", "DOMMx Final Assessment Report")

        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Scope: {meta.get('scope', '')}")
        doc.add_paragraph(f"Generated: {meta.get('generated_at', '')}")
        doc.add_paragraph(f"Engine version: {REPORT_ENGINE_VERSION}")

        doc.add_page_break()

    # ==========================================================
    # TOC
    # ==========================================================

    def _render_toc(self, doc: Document):

        doc.add_heading("Table of Contents", level=1)

        p = doc.add_paragraph()
        run = p.add_run()

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)

        doc.add_page_break()

    # ==========================================================
    # RESULTS
    # ==========================================================

    def _render_results_section(self, doc: Document, report_json: Dict[str, Any]):

        domains = report_json.get("domains", []) or []

        doc.add_heading("1. Consolidated Results", level=1)

        table = doc.add_table(rows=1, cols=5)

        hdr = table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "Domain"
        hdr[2].text = "Name"
        hdr[3].text = "Score"
        hdr[4].text = "Level"

        total = 0
        count = 0

        for idx, d in enumerate(domains, start=1):

            level = d.get("level", 0)

            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = str(d.get("acronym", ""))
            row[2].text = str(d.get("name", ""))
            row[3].text = str(level)
            row[4].text = self._level_label(level)

            if isinstance(level, (int, float)):
                total += level
                count += 1

        if count > 0:
            avg = round(total / count, 2)
            doc.add_paragraph(f"\nOverall Average Maturity: {avg}")

        doc.add_page_break()

    # ==========================================================
    # DEMO SECTION
    # ==========================================================

    def _render_demo_section(self, doc: Document, report_json: Dict[str, Any]):

        domains = report_json.get("domains", []) or []

        doc.add_heading("2. Recommended Structural Documentation", level=1)

        for domain in domains:

            acronym = domain.get("acronym")
            name = domain.get("name")
            level = domain.get("level")
            action_code = self._extract_action_code(domain)

            doc.add_page_break()

            doc.add_heading(f"{acronym} - {name}", level=2)
            doc.add_paragraph(f"Maturity Level: {level}")
            doc.add_paragraph(f"Recommended Action: {action_code}")

            demo_texts = self._load_demo_texts(acronym, action_code)

            if not demo_texts:
                doc.add_paragraph("No structural demo content found.")
                continue

            for demo in demo_texts:
                self._write_demo_text(doc, demo)

    # ==========================================================
    # LOAD PATCHED FILE
    # ==========================================================

    def _load_demo_texts(
        self,
        acronym: str,
        action_code: str
    ) -> List[str]:

        if not acronym or not action_code:
            return []

        demo_path = (
            self.base_dir
            / "data"
            / "domains"
            / self.catalog_lang
            / f"{acronym}_theory_demo_output_PATCHED.json"
        )

        if not demo_path.exists():
            return []

        with open(demo_path, "r", encoding="utf-8") as f:
            payload = json.load(f)

        return [
            entry.get("demo")
            for entry in payload
            if entry.get("action_code") == action_code
            and entry.get("demo")
        ]

    # ==========================================================
    # WRITE STRUCTURED DEMO
    # ==========================================================

    def _write_demo_text(self, doc: Document, demo_text: str):

        for line in demo_text.split("\n"):

            stripped = line.strip()

            if not stripped:
                doc.add_paragraph("")
                continue

            if self._is_main_section(stripped):
                doc.add_heading(stripped, level=3)
            elif self._is_sub_section(stripped):
                doc.add_heading(stripped, level=4)
            elif stripped.startswith("-"):
                doc.add_paragraph(stripped, style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    def _is_main_section(self, text: str) -> bool:
        return text[:2].isdigit() and "." in text[:4]

    def _is_sub_section(self, text: str) -> bool:
        parts = text.split(".")
        return len(parts) >= 3 and parts[0].isdigit()

    # ==========================================================
    # ACTION CODE
    # ==========================================================

    def _extract_action_code(self, domain_model: Dict[str, Any]) -> str:

        for e in domain_model.get("level_entries", []) or []:
            code = str(e.get("Action_Code") or "").strip()
            if code:
                return code

        for a in domain_model.get("actions", []) or []:
            code = str(a.get("action_code") or "").strip()
            if code:
                return code

        return ""

    # ==========================================================
    # LEVEL LABEL
    # ==========================================================

    def _level_label(self, level):

        if level is None:
            return "Not Assessed"

        try:
            level_int = int(level)
        except (TypeError, ValueError):
            return str(level)

        labels = {
            0: "Initial",
            1: "Ad Hoc",
            2: "Repeatable",
            3: "Defined",
            4: "Managed",
            5: "Optimized",
        }

        return labels.get(level_int, str(level_int))