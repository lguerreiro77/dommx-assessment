# root/core/ai_report_service.py

from __future__ import annotations

import os
import re
import json
import math
import hashlib
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from openai import OpenAI
import streamlit as st
import yaml

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class DomainMeta:
    domain_key: str                 # domain_0, domain_1...
    acronym: str                    # DG
    name: str                       # Data Governance
    domain_id: str                  # as in flow.yaml domain_id
    sequence: str                   # flow sequence
    dependence: List[str]           # flow dependence list
    qtext: Dict[str, str]           # qid lower -> question text


@dataclass
class DomainScore:
    domain_key: str
    acronym: str
    name: str
    # [(qid, question_text, score_float)]
    question_scores: List[Tuple[str, str, float]]
    avg_raw: float
    avg_floor: int


@dataclass
class DependencyIssue:
    domain_acronym: str
    reference_acronym: str
    dependency_broken: bool
    severity_rationale: str
    scenarios: Dict[str, Any]


@dataclass
class UserMeta:
    user_id: str    
    full_name: str
    email: str


class AIReportService:
    """
    Enterprise report generator aligned with export_service.py logic:
      - Domain/question mapping from FileSystem_Setup.yaml -> flow.yaml + orchestration.yaml (execution_request) + decision_tree.yaml
      - Results scope strictly from 'results' table (dbassessment.xlsx equivalent)

    Scope rules:
      - Non-admin: report for (project_id + user_id)
      - Admin: consolidated report for (project_id + all users), averaging answers per question

    Outputs:
      - DOCX
      - PDF (simple corporate)
    """
    
###----------------------------------------------------------------------------------------------------------------------------------------------------            
    
    def _load_dependency_inconsistency_theory(self, language: str):

        import json

        lang = (language or "us").lower()

        path = (
            self.base_dir
            / "data"
            / "domains"
            / lang
            / "Dependencies_inconsistencies_theory_cluster_output.json"
        )

        if not path.exists():
            return []

        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        if isinstance(data, dict):
            return data.get("inconsistencies", [])

        if isinstance(data, list):
            return data

        return []


    def _detect_structural_dependency_inconsistencies(
        self,
        domain_metas: Dict[str, DomainMeta],
        scores: List[DomainScore]
    ) -> List[Dict[str, Any]]:

        # usa avg_floor para regra estrutural do relatório
        score_by_acr = {s.acronym: s.avg_floor for s in scores}

        # domain_id → acronym lookup
        id_to_acr = {str(m.domain_id): m.acronym for m in domain_metas.values()}

        inconsistencies: List[Dict[str, Any]] = []

        for m in domain_metas.values():

            child_acr = m.acronym
            child_score = score_by_acr.get(child_acr)

            if child_score is None:
                continue

            deps = m.dependence or []
            deps_only = deps[:-1] if len(deps) >= 2 else deps  # regra do flow

            for dep_id in deps_only:

                parent_acr = id_to_acr.get(str(dep_id))
                if not parent_acr:
                    continue

                parent_score = score_by_acr.get(parent_acr)
                if parent_score is None:
                    continue

                gap = child_score - parent_score

                # regra estrutural: só consideramos ruptura quando gap >= 1
                if gap < 1:
                    continue

                # evitar duplicação
                if any(x["child"] == child_acr and x["parent"] == parent_acr for x in inconsistencies):
                    continue

                inconsistencies.append(
                    {
                        "child": child_acr,
                        "parent": parent_acr,
                        "child_score": child_score,
                        "parent_score": parent_score,
                        "gap": gap
                    }
                )

        # ordem consistente: segue dependências do fluxo (child depois parent)
        inconsistencies = sorted(inconsistencies, key=lambda x: (x["child"], x["parent"]))
        return inconsistencies


    def _render_dependency_breaks(
        self,
        doc: Document,
        inconsistencies: List[Dict[str, Any]],
        theory_data: List[Dict[str, Any]],
        language: str,
        acr_to_name: Optional[Dict[str, str]] = None
    ):

        import re

        if not inconsistencies:
            self._add_paragraph(doc, self._t("no_dependency_breaks", language))
            return

        acr_to_name = acr_to_name or {}

        # index rápido (child, parent) -> item do theory json
        theory_map: Dict[tuple, Dict[str, Any]] = {}
        for t in (theory_data or []):
            c = (t.get("domain_acronym") or "").strip()
            p = (t.get("reference_acronym") or "").strip()
            if c and p:
                theory_map.setdefault((c,p), t)

        dep_index = 1
        
        
        seen = set()

        for inc in inconsistencies:

            child = inc.get("child")
            parent = inc.get("parent")
            gap = inc.get("gap", 0)

            # evita duplicação do mesmo par
            key = (child, parent)
            if key in seen:
                continue
            seen.add(key)

            child_name = acr_to_name.get(child, "")
            parent_name = acr_to_name.get(parent, "")

            # 2.2.1, 2.2.2, ...
            heading = f"2.2.{dep_index} - {child} · {child_name} {self._t('depends_on', language)} {parent} · {parent_name}".strip()
            self._add_heading(doc, heading, level=3)
            dep_index += 1

            th = theory_map.get((child, parent)) or theory_map.get((parent, child))

            if th:

                sev = (
                    (th.get("Structural Severity Classification") or {})
                    .get("severity_rationale")
                )

                if sev and len(sev) < 250:
                    self._add_paragraph(
                        doc,
                        f"{self._t('severity_rationale_label', language)} {sev}"
                    )

                scenarios = th.get("scenarios") or {}

                scenario = scenarios.get("reference_inferior")

                if gap < 1:
                    scenario = scenarios.get("reference_superior")

                if scenario is None:
                    scenario = scenarios.get("reference_not_evaluated")

                txt = scenario.get("analysis_text") if scenario else None

                if txt and isinstance(txt, str):

                    txt = re.sub(
                        r"(?m)^\s*#{1,6}\s*\d+\)\s*(.+)$",
                        r"• \1",
                        txt
                    )

                    txt = re.sub(r"(?m)^#{1,6}\s*", "", txt)
                    
                    txt = txt.replace("\r", "")                  
                    
                    sections = re.split(r"(?=^\s*• )", txt, flags=re.MULTILINE)

                    for section in sections:
                        clean = section.strip()
                        if clean:
                            self._add_paragraph(doc, clean)

            else:

                fallback = (
                    f"{child} is more mature than its structural dependency {parent}. "
                    f"This usually means the dependent domain is advancing without the needed foundation in the dependency."
                )

                self._add_paragraph(doc, fallback)

            doc.add_paragraph("")
   
            
            
###----------------------------------------------------------------------------------------------------------------------------------------------------            
            
    def _normalize_maturity_labels_in_demo(
        self,
        text: str,
        tree_data: Dict[str, Any]
    ) -> str:
        """
        Substitui qualquer ocorrência de:
        'Nível X - Algo'
        'Pontuação X - Algo'
        Pelo label oficial definido no maturity_scale do decision_tree.
        """

        if not text or not isinstance(tree_data, dict):
            return text

        maturity_scale = tree_data.get("maturity_scale") or {}
        if not isinstance(maturity_scale, dict):
            return text

        # construir mapa oficial
        official = {}
        for k, v in maturity_scale.items():
            try:
                level = int(k)
            except Exception:
                continue

            label = str(v.get("maturity_level") or "").strip()
            if label:
                official[level] = label

        if not official:
            return text

        import re

        def replace(match):
            level = int(match.group(1))
            if level in official:
                return f"Nível {level} - {official[level]}"
            return match.group(0)

        # captura Nível 3 - Algo
        pattern1 = re.compile(r"Nível\s+(\d+)\s*-\s*([^\n]+)")
        text = pattern1.sub(replace, text)

        # captura Pontuação 3 - Algo
        pattern2 = re.compile(r"Pontuação\s+(\d+)\s*-\s*([^\n]+)")
        text = pattern2.sub(replace, text)

        return text
    
    # =========================================================
    # Radar Analysis Helpers (DOMMx)
    # =========================================================

    client = OpenAI()

    def _polish_text_with_ai(self, analysis: Dict[str, Any], language: str) -> str:

        key = hashlib.md5((json.dumps(analysis) + language).encode()).hexdigest()

        if not hasattr(self, "_ai_cache"):
            self._ai_cache = {}

        if key in self._ai_cache:
            return self._ai_cache[key]
            
        
        scores = "\n".join(
            f"{d}: {s}"
            for d, s in analysis["sorted_domains"]
        )        
        

        prompt = f"""
    Write a concise, professional interpretation of a data governance maturity radar chart.

    Language: {language}

    Inputs (do not repeat numbers verbatim in the output):
    - maturity_index: {analysis['maturity_index']}
    - balance_index: {analysis['balance_index']}
    - risk_index: {analysis['risk_index']}
    - domain_scores: {analysis['sorted_domains']}

    Rules:
    1) No title. No asterisks. No headings.    
    2) Output only analytical narrative text. Do not list scores, indices, or domain numbers.
    3) Use 2–3 short paragraphs. Bullet points only if strictly necessary.
    4) Disparity logic:
       - Only describe “major disparity” if any domain differs by more than 1.0 point from the average score.
       - If all domains are within 0.5 points of each other, explicitly state the profile is balanced.
    5) Priority logic:
       - Priority domains are those below the average score.
       - If none are below average, say priorities are incremental and focus on consistency and scaling.
    6) Risk interpretation:
       - Interpret risk_index only as “exposure driven by the lowest-maturity domains” (do not assume security/compliance).
    7) Recommendations:
       - When suggesting priorities, briefly explain how to improve in generic governance terms (process, roles, standards, monitoring), without mentioning specific tools or technologies.
    8) Always finish the last sentence.
    9) Keep it under 250 words.    
        """

        try:

            resp = self.client.chat.completions.create(
                model="gpt-4o-mini",
                temperature=0.7,
                max_tokens=500,
                messages=[
                    {"role": "system", "content": "You are an expert in data governance maturity assessments."},
                    {"role": "user", "content": prompt}
                ]
            )

            text = resp.choices[0].message.content.strip()

            self._ai_cache[key] = text

            return text

        except Exception:
            return ""
        
    def _compute_radar_analysis(self, domains_scores: Dict[str, float], max_score: float = 5.0):

        from statistics import mean, pstdev

        scores = list(domains_scores.values())
        n = len(scores)

        avg = mean(scores)
        std = pstdev(scores) if n > 1 else 0

        min_score = min(scores)
        max_score_obs = max(scores)

        maturity_index = avg / max_score

        balance_index = 1 - ((max_score_obs - min_score) / max_score)

        risk_index = 1 - (min_score / max_score)

        leaders = []
        critical = []

        for d, s in domains_scores.items():

            if s >= avg + std:
                leaders.append(d)

            elif s <= avg - std:
                critical.append(d)

        sorted_domains = sorted(
            domains_scores.items(),
            key=lambda x: x[1]
        )
        

        return {
            "domain_count": n,
            "maturity_index": round(maturity_index, 2),
            "balance_index": round(balance_index, 2),
            "risk_index": round(risk_index, 2),
            "leaders": leaders,
            "critical": critical,
            "sorted_domains": sorted_domains
        }


    def _generate_radar_analysis_text(self, analysis: Dict[str, Any], language: str) -> str:

        lang = (language or "us").lower()

        # report_text por idioma (já carregado por _load_report_texts)
        texts = self._report_texts.get(lang) or self._report_texts.get("us") or {}

        # templates (carregados fora do bloco report_text)
        templates = getattr(self, "_analysis_templates", None) or []
        if not templates:
            return ""

        import random
        template = random.choice(templates)

        leaders = ", ".join(analysis.get("leaders") or []) if analysis.get("leaders") else texts.get("tag_no_leader", "")
        critical = ", ".join(analysis.get("critical") or []) if analysis.get("critical") else texts.get("tag_no_critical", "")

        maturity_index = float(analysis.get("maturity_index", 0.0))
        balance_index = float(analysis.get("balance_index", 0.0))
        risk_index = float(analysis.get("risk_index", 0.0))

        # -------- labels via tags (sem texto hardcoded) --------
        def _maturity_label(v: float) -> str:
            # v esperado 0..5
            lvl = int(round(v))
            lvl = max(0, min(5, lvl))
            if lvl <= 1:
                return texts.get("tag_maturity_initial", "")
            if lvl == 2:
                return texts.get("tag_maturity_emerging", "")
            if lvl == 3:
                return texts.get("tag_maturity_established", "")
            if lvl == 4:
                return texts.get("tag_maturity_managed", "")
            return texts.get("tag_maturity_optimized", "")

        def _balance_label(v: float) -> str:
            # v típico 0..1 (quanto maior, mais equilibrado)
            if v >= 0.80:
                return texts.get("tag_balance_high", "")
            if v >= 0.60:
                return texts.get("tag_balance_moderate", "")
            if v >= 0.40:
                return texts.get("tag_balance_low", "")
            return texts.get("tag_balance_fragmented", "")

        def _risk_label(v: float) -> str:
            # v típico 0..1 (quanto maior, mais risco)
            if v <= 0.25:
                return texts.get("tag_risk_low", "")
            if v <= 0.50:
                return texts.get("tag_risk_moderate", "")
            if v <= 0.75:
                return texts.get("tag_risk_high", "")
            return texts.get("tag_risk_critical", "")

        context = {
            "maturity_index": maturity_index,
            "balance_index": balance_index,
            "risk_index": risk_index,
            "domain_count": analysis.get("domain_count", 0),
            "leaders": leaders,
            "critical": critical,
            "maturity_label": _maturity_label(maturity_index),
            "balance_label": _balance_label(balance_index),
            "risk_label": _risk_label(risk_index),
        }

        paragraphs = []

        # IMPORTANTE:
        # Os templates usam {tag_*} e também {maturity_index}, etc.
        # Então a gente passa texts + context no format.
        paragraphs.append(template.get("p1", "").format(**texts, **context))
        paragraphs.append(template.get("p2", "").format(**texts, **context))
        paragraphs.append(template.get("p3", "").format(**texts, **context))

        # parágrafo por domínio (funciona com 3 ou 12 sem mudar nada)
        for domain, score in (analysis.get("sorted_domains") or []):
            if domain in (analysis.get("critical") or []):
                status = texts.get("tag_domain_status_critical", "")
            elif domain in (analysis.get("leaders") or []):
                status = texts.get("tag_domain_status_leader", "")
            else:
                status = texts.get("tag_domain_status_aligned", "")

            p = texts.get("domain_paragraph", "{domain} {score} {status}")
            paragraphs.append(
                p.format(
                    domain=domain,
                    score=score,
                    status=status
                )
            )

        # limpa vazios (evita parágrafo em branco “mudo”)
        paragraphs = [p for p in paragraphs if (p or "").strip()]

        return "\n\n".join(paragraphs)
        
    # =========================================================
    
    def _add_radar_chart(self, doc: Document, scores: List[DomainScore], language: str):

        if not scores:
            return

        # radar precisa pelo menos 2 eixos para ficar decente
        if len(scores) < 2:
            return

        import numpy as np
        import matplotlib.pyplot as plt
        from io import BytesIO
        from docx.shared import Inches

        labels = [s.acronym for s in scores]
        values = [float(s.avg_raw) for s in scores]

        angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
        values = values + values[:1]
        angles = angles + angles[:1]

        fig = plt.figure(figsize=(4.5, 4.5))
        ax = fig.add_subplot(111, polar=True)

        ax.plot(angles, values, linewidth=2)
        ax.fill(angles, values, alpha=0.20)

        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(labels)

        ax.set_ylim(0, 5)
        ax.set_yticks([0, 1, 2, 3, 4, 5])

        ax.set_title("Domain Maturity Radar", pad=20)

        buf = BytesIO()
        plt.tight_layout()
        plt.savefig(buf, format="png", dpi=200)
        plt.close(fig)
        buf.seek(0)
        
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(buf, width=Inches(3.75))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # =====================================================
        # Radar Analysis Text (DOMMx)
        # =====================================================

        domains_scores = {}

        for s in scores:
            try:
                domains_scores[s.acronym] = float(s.avg_raw)
            except Exception:
                pass

        if len(domains_scores) >= 2:

            analysis = self._compute_radar_analysis(domains_scores)
                        
            cache_key = hashlib.md5((json.dumps(analysis) + language).encode()).hexdigest()

            if not hasattr(self, "_ai_cache"):
                self._ai_cache = {}

            if cache_key in self._ai_cache:
                analysis_text = self._ai_cache[cache_key]
            else:
                analysis_text = self._polish_text_with_ai(
                    analysis,
                    language
                )

            if analysis_text:

                #doc.add_paragraph("")  # 1 linha após a figura

                title = self._t("tag_analysis_radar", language)

                p_title = doc.add_paragraph()
                run_title = p_title.add_run(title)
                run_title.bold = True

                #doc.add_paragraph("")  # 1 linha após o título

                for block in analysis_text.split("\n\n"):
                    doc.add_paragraph(block)

                #doc.add_paragraph("")
        

    def _likert_label(self, grade: int, lang: str) -> str:
        texts = self._report_texts or {}
        lang = (lang or "us").lower()

        if lang in texts and "likert" in texts[lang]:
            return texts[lang]["likert"].get(int(grade), "")

        if "us" in texts and "likert" in texts["us"]:
            return texts["us"]["likert"].get(int(grade), "")

        return str(grade)
    
    def _load_report_texts(self) -> Dict[str, Dict[str, str]]:
        """
        Loads i18n report texts from:
            root/data/general/report_texts.yaml
        Fallback: empty dict
        """
        path = self.base_dir / "data" / "general" / "report_texts.yaml"

        if not path.exists():
            return {}

        try:
            data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
            return data.get("report_text", {})
        except Exception:
            return {}
            
    def _load_analysis_templates(self) -> List[Dict[str, str]]:
        """
        Loads radar analysis templates from:
            root/data/general/report_texts.yaml
        Returns a list of dicts: [{"p1": "...", "p2": "...", "p3": "..."}, ...]
        Fallback: empty list
        """
        path = self.base_dir / "data" / "general" / "report_texts.yaml"

        if not path.exists():
            return []

        try:
            data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
            templates = data.get("analysis_templates", []) or []
            # garante formato esperado
            if isinstance(templates, list):
                return [t for t in templates if isinstance(t, dict)]
            return []
        except Exception:
            return []   
    
            

    def _t(self, key: str, lang: str) -> str:
        lang = (lang or "us").lower()

        texts = self._report_texts or {}

        # idioma específico
        if lang in texts and key in texts[lang]:
            return texts[lang][key]

        # fallback para us
        if "us" in texts and key in texts["us"]:
            return texts["us"][key]

        # fallback final
        return key
        
    
    def __init__(self, base_dir: str, repo: Any):
        self.base_dir = Path(base_dir)
        self.repo = repo        

        # Load report texts from YAML
        self._report_texts = self._load_report_texts()
        self._analysis_templates = self._load_analysis_templates()
                
        
    def _force_update_fields_on_open(self, doc: Document):
        settings = doc.settings._element
        update = OxmlElement('w:updateFields')
        update.set(qn('w:val'), 'true')
        settings.append(update)
        
        
    def _override_domain_definition_kv(
        self,
        block: Dict[str, Any],
        domain_acronym: str,
        domain_name: str
    ):
        """
        Se o KV for 'Definição do Domínio',
        substitui o valor por 'DG - Governança de Dados'
        """
        key_norm = (block.get("key") or "").strip().lower()

        if key_norm in ["definição do domínio", "definicao do dominio"]:
            # regra pedida: código do domínio + traço + nome do domínio
            block["text"] = f"{domain_acronym} - {domain_name}"

    # =========================================================
    # Public API
    # =========================================================

    def generate_report_docx(
        self,
        project_id: str,
        user_id: str,
        is_admin: bool,
        language: str,
        force_regen: bool = False
    ) -> str:
        out_dir, meta = self._prepare_cache(project_id, user_id, is_admin, language, force_regen)
        docx_path = out_dir / meta["docx_name"]

        if docx_path.exists() and not force_regen:
            #force_regen=True
            return str(docx_path)

        # Build mappings like export_service (filesystem -> flow + orchestration -> decision_tree)
        mapping = self._load_project_mapping(project_id, language=language)
        domain_metas: Dict[str, DomainMeta] = mapping["domain_metas"]
        flow = mapping["flow"]
        orch = mapping["orch"]
        resolved_lang = (language or "us").lower()
        
        T = lambda k: self._t(k, resolved_lang)

        # Extract answers from results table (strict scope)
        answers_by_domain, included_users = self._load_answers(project_id, user_id, is_admin)

        # Build scores strictly for domains present in results/answers and ordered by execution_request/domain_key
        scores = self._compute_scores(domain_metas, answers_by_domain)

        # Dependencies issues knowledge base
        deps_issues = self._load_dependency_issues()

        # Project/User metadata for cover
        project_name = self._get_project_name(project_id)
        user_meta = self._get_user_meta(user_id) if not is_admin else None

        doc = Document()
        self._setup_doc_styles(doc)
        self._force_update_fields_on_open(doc)
        
        # Cover
        self._add_cover(
            doc=doc,
            project_id=project_id,
            project_name=project_name,
            user_meta=user_meta,
            is_admin=is_admin,
            included_users=included_users,
            language=resolved_lang
        )

        # -------------------------------------------------
        # Cover Purpose – DOMMx
        # -------------------------------------------------
        purpose_text = T("cover_purpose_dommmx")
        if purpose_text and purpose_text != "cover_purpose_dommmx":
            doc.add_paragraph("")
            self._add_paragraph(doc, purpose_text)

        # -------------------------------------------------
        # How to Use This Report
        # -------------------------------------------------
        how_text = T("how_to_use_report")

        if how_text and how_text != "how_to_use_report":

            self._add_page_break(doc)

            title_txt = T("how_to_use_report_title")
            if not title_txt or title_txt == "how_to_use_report_title":
                title_txt = "How to Use This Report"

            # Título visual sem entrar no TOC
            p = doc.add_paragraph()
            run = p.add_run(title_txt)
            run.bold = True
            run.font.size = Pt(18)

            doc.add_paragraph("")
            self._add_paragraph(doc, how_text)

        # -------------------------------------------------
        # Table of Contents (visual only, not structural)
        # -------------------------------------------------
        self._add_page_break(doc)

        toc_title = T("toc")
        if not toc_title or toc_title == "toc":
            toc_title = "Table of Contents"

        p = doc.add_paragraph()
        run = p.add_run(toc_title)
        run.bold = True
        run.font.size = Pt(20)

        doc.add_paragraph("")
        self._add_toc_field(doc)
        

        # 1) Results        
        self._add_page_break(doc)
        self._add_heading(doc, T("section_1"), level=1)
        self._add_paragraph(doc, T("section_1_intro_1"))
        self._add_paragraph(doc, T("section_1_intro_2"))
        self._add_results_section(doc, scores, is_admin=is_admin, language=resolved_lang)

        # 2) Dependencies
        self._add_page_break(doc)
        self._add_heading(doc, T("section_2"), level=1)
        self._add_paragraph(doc, T("section_2_intro"))
        
        # ---------------------------------------------------------
        # 2.1 Declared Dependencies
        # ---------------------------------------------------------
        self._add_paragraph(doc, "")
        self._add_heading(doc, T("declared_dependencies_heading"), level=2)

        score_by_acr = {s.acronym: s.avg_floor for s in scores}
        in_scope_acr = {s.acronym for s in scores}

        for m in domain_metas.values():

            if m.acronym not in in_scope_acr:
                continue

            self._add_heading(doc, f"{m.acronym} · {m.name}", level=3)

            dep = m.dependence or []

            if not dep:
                self._add_paragraph(doc, T("no_dependencies"))
                continue

            deps_only = dep[:-1] if len(dep) >= 2 else dep

            if not deps_only:
                self._add_paragraph(doc, T("no_dependencies"))
                continue

            self._add_paragraph(doc, T("declared_dependencies_label"))

            for x in deps_only:

                for mm in domain_metas.values():

                    if str(mm.domain_id) == str(x):

                        dep_acr = mm.acronym
                        dep_name = mm.name

                        dep_grade = score_by_acr.get(dep_acr)

                        if dep_grade is None:

                            doc.add_paragraph(
                                f"{dep_acr} · {dep_name} ({T('not_evaluated_scope')})",
                                style="List Bullet"
                            )

                        else:

                            doc.add_paragraph(
                                f"{dep_acr} · {dep_name} "
                                f"(grade {dep_grade} · {self._likert_label(dep_grade, resolved_lang)})",
                                style="List Bullet"
                            )
        
        #2.2
        
        doc.add_paragraph("")
        self._add_heading(doc, self._t("detected_breaks", resolved_lang), level=2)
        
        theory_data = self._load_dependency_inconsistency_theory(resolved_lang)

        inconsistencies = self._detect_structural_dependency_inconsistencies(
            domain_metas,
            scores
        )
        
        doc.add_paragraph("")
         
        self._render_dependency_breaks(
            doc,
            inconsistencies,
            theory_data,
            resolved_lang
        )

        # 3) Blueprint
        self._add_page_break(doc)
        self._add_heading(doc, T("section_3"), level=1)
        self._add_paragraph(doc, T("section_3_intro"))
        
        self._add_blueprint_section(
            doc=doc,
            project_id=project_id,
            scores=scores,
            domain_metas=domain_metas,
            flow=flow,
            language=resolved_lang
        )

        # 4) References
        self._add_page_break(doc)
        self._add_heading(doc, T("section_4"), level=1)
        self._add_references_section(doc, resolved_lang)

        doc.save(str(docx_path))

        # PDF
        self._export_pdf(
            pdf_path=str(out_dir / meta["pdf_name"]),
            project_id=project_id,
            project_name=project_name,
            user_meta=user_meta,
            is_admin=is_admin,
            included_users=included_users,
            language=resolved_lang,
            scores=scores,
            domain_metas=domain_metas,
            deps_issues=deps_issues,
        )

        (out_dir / "report.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

        return str(docx_path)

    def generate_report_pdf(
        self,
        project_id: str,
        user_id: str,
        is_admin: bool,
        language: str,
        force_regen: bool = False
    ) -> str:
        out_dir, meta = self._prepare_cache(project_id, user_id, is_admin, language, force_regen)
        pdf_path = out_dir / meta["pdf_name"]

        if pdf_path.exists() and not force_regen:
            return str(pdf_path)

        _ = self.generate_report_docx(project_id, user_id, is_admin, language, force_regen=force_regen)
        if not pdf_path.exists():
            raise RuntimeError("PDF export failed to generate.")

        return str(pdf_path)

    # =========================================================
    # Cache
    # =========================================================

    def _prepare_cache(self, project_id: str, user_id: str, is_admin: bool, language: str, force_regen: bool):
        language = (language or "us").strip()

        answers_by_domain, included_users = self._load_answers(project_id, user_id, is_admin)

        mapping = self._load_project_mapping(project_id, language=language)
        flow_hash = self._hash_text(yaml.safe_dump(mapping["flow"], sort_keys=True))
        orch_hash = self._hash_text(yaml.safe_dump(mapping["orch"], sort_keys=True))

        fingerprint = self._fingerprint({
            "schema": "report_v2_enterprise",
            "project_id": project_id,
            "user_id": user_id,
            "is_admin": bool(is_admin),
            "language": (language or "us").strip().lower(),
            "answers": answers_by_domain,
            "included_users": [u.user_id for u in included_users],
            "flow_hash": flow_hash,
            "orch_hash": orch_hash,
        })

        reports_root = self.base_dir / "data" / "reports"

        project_name = self._get_project_name(project_id)
        safe_project = re.sub(r'[^A-Za-z0-9_]+', '_', project_name or project_id)
        safe_user = "ALL_USERS" if is_admin else re.sub(r'[^A-Za-z0-9_]+', '_', user_id or "")

        out_dir = reports_root / safe_project / safe_user
        
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")

        # 🔥 CALCULAR VARIÁVEIS ANTES DO DICT
        safe_project = re.sub(r'[^A-Za-z0-9_]+', '_', project_name or "")
        safe_user = "ALL_USERS" if is_admin else re.sub(r'[^A-Za-z0-9_]+', '_', user_id or "")
        seq = datetime.utcnow().strftime("%Y%m%d_%H%M%S")

        meta = {
            "schema": "report_v2_enterprise",
            "project_id": project_id,
            "user_id": user_id,
            "is_admin": bool(is_admin),
            "language": (language or "us").strip().lower(),
            "generated_utc": ts,
            "docx_name": f"{safe_project}_{safe_user}_{seq}.docx",
            "pdf_name": f"{safe_project}_{safe_user}_{seq}.pdf",
            "fingerprint": fingerprint,
        }

        if force_regen:
            for fn in ("report.json", meta["docx_name"], meta["pdf_name"]):
                p = out_dir / fn
                if p.exists():
                    try:
                        p.unlink()
                    except Exception:
                        pass

        return out_dir, meta

    def _hash_text(self, text: str) -> str:
        return hashlib.sha256((text or "").encode("utf-8")).hexdigest()[:16]

    def _fingerprint(self, payload: Any) -> str:
        raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)
        return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]

    # =========================================================
    # Export-service-aligned mapping
    # =========================================================

    def _load_project_mapping(self, project_id: str, language: str) -> Dict[str, Any]:

        fs_path = self._resolve_filesystem_setup(project_id)
        fs_setup = self._safe_load_yaml(fs_path) or {}
        orch_cfg = (fs_setup.get("orchestrator_config") or {})

        flow_rel = orch_cfg.get("main_flow", "flow.yaml")
        orch_rel = orch_cfg.get("main_orchestration", "default_execution.yaml")

        flow_path = self._resolve_project_file(project_id, flow_rel)
        orch_path = self._resolve_project_file(project_id, orch_rel)

        flow = self._safe_load_yaml(flow_path) or {}
        orch = self._safe_load_yaml(orch_path) or {}

        req_list = orch.get("execution_request", []) or []
        domain_flow = flow.get("Domain_flow", []) or []

        # 🔥 idioma estrutural do projeto (para decision_tree)
        project_lang = str(orch.get("language") or "us").strip().lower()
        if not project_lang:
            project_lang = "us"

        domain_metas: Dict[str, DomainMeta] = {}

        for idx, req in enumerate(req_list):
            dom_id = req.get("domain")
            dom_meta = next(
                (d for d in domain_flow if str(d.get("domain_id")) == str(dom_id)),
                None
            ) or {}

            acronym = (dom_meta.get("acronym") or str(dom_id) or f"domain_{idx}").strip()
            name = (dom_meta.get("name") or "").strip()
            sequence = str(dom_meta.get("sequence") or "").strip()
            dependence = dom_meta.get("dependence") or []
            dependence = [str(x) for x in dependence if x is not None]

            files = dom_meta.get("files") or {}
            decision_tree = files.get("decision_tree")

            qtext_map: Dict[str, str] = {}

            if decision_tree:
                tree_path = (
                    self.base_dir
                    / "data"
                    / "domains"
                    / "Language"
                    / project_lang
                    / decision_tree
                )

                tree_data = self._safe_load_yaml(tree_path) or {}
                questions = tree_data.get("questions", {}) or {}

                for qid, qinfo in questions.items():
                    qid_str = str(qid).strip().lower()
                    qtext = (qinfo.get("question") or qinfo.get("text") or "").strip()
                    if qid_str:
                        qtext_map[qid_str] = qtext

            domain_key = f"domain_{idx}"

            domain_metas[domain_key] = DomainMeta(
                domain_key=domain_key,
                acronym=acronym,
                name=name,
                domain_id=str(dom_id) if dom_id is not None else "",
                sequence=sequence,
                dependence=dependence,
                qtext=qtext_map,
            )

        return {
            "fs_path": str(fs_path),
            "flow_path": str(flow_path),
            "orch_path": str(orch_path),
            "flow": flow,
            "orch": orch,
            "project_lang": project_lang,   # 🔥 apenas estrutural
            "domain_metas": domain_metas,
        }

    def _resolve_filesystem_setup(self, project_id: str) -> Path:
        # Prefer project General/FileSystem_Setup.yaml, then project root, then BASE_DIR root
        p1 = self.base_dir / "data" / "projects" / str(project_id) / "General" / "FileSystem_Setup.yaml"
        if p1.exists():
            return p1
        p2 = self.base_dir / "data" / "projects" / str(project_id) / "FileSystem_Setup.yaml"
        if p2.exists():
            return p2
        p3 = self.base_dir / "FileSystem_Setup.yaml"
        if p3.exists():
            return p3
        raise FileNotFoundError("FileSystem_Setup.yaml not found (project or base).")

    def _resolve_project_file(self, project_id: str, rel_path: str) -> Path:
        """
        Resolves a yaml path referenced by filesystem/orchestrator.
        Tries:
          - project General/<rel_path>
          - project root/<rel_path>
          - BASE_DIR/<rel_path>
        """
        rel = str(rel_path or "").strip()
        if not rel:
            raise FileNotFoundError("Empty config path.")

        # If absolute, use directly
        rp = Path(rel)
        if rp.is_absolute() and rp.exists():
            return rp

        candidates = [
            self.base_dir / "data" / "projects" / str(project_id) / "General" / rel,
            self.base_dir / "data" / "projects" / str(project_id) / rel,
            self.base_dir / rel,
        ]
        for c in candidates:
            if c.exists():
                return c

        # last fallback: try lower/upper for "general/General"
        candidates2 = [
            self.base_dir / "data" / "projects" / str(project_id) / "general" / rel,
        ]
        for c in candidates2:
            if c.exists():
                return c

        raise FileNotFoundError(f"Project file not found: {rel}")

    def _safe_load_yaml(self, path: Path) -> Optional[Dict[str, Any]]:
        try:
            return yaml.safe_load(path.read_text(encoding="utf-8"))
        except Exception:
            return None

    def _safe_load_json(self, path: Path) -> Optional[Any]:
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return None

    # =========================================================
    # Results extraction (strict scope)
    # =========================================================

    def _load_answers(
        self,
        project_id: str,
        user_id: str,
        is_admin: bool
    ) -> Tuple[Dict[str, Dict[str, float]], List[UserMeta]]:
        """
        Returns:
          answers_by_domain: {domain_key: {qid: score_float}}
          included_users: users included in report scope

        Scope:
          - Non-admin: only (project_id, user_id)
          - Admin: all users for project_id, averaged per question
        """
        results = self.repo.fetch_all("results") or []
        if not results:
            return {}, []

        # filter rows by scope
        scoped_rows = []
        for r in results:
            if str(r.get("project_id")) != str(project_id):
                continue
            if not is_admin and str(r.get("user_id")) != str(user_id):
                continue
            scoped_rows.append(r)

        if not scoped_rows:
            return {}, []

        # included users metadata
        included_user_ids = sorted({str(r.get("user_id") or "").strip() for r in scoped_rows if str(r.get("user_id") or "").strip()})
        included_users = [self._get_user_meta(uid) for uid in included_user_ids]

        # decrypt and aggregate
        from auth.crypto_service import decrypt_text

        # aggregated[domain_key][qid] -> list of float scores
        aggregated: Dict[str, Dict[str, List[float]]] = {}

        for r in scoped_rows:
            enc = r.get("answers_json_encrypted")
            if not enc:
                continue

            try:
                decrypted = decrypt_text(enc)
                payload = json.loads(decrypted)
            except Exception:
                continue

            answers = payload.get("answers") if isinstance(payload, dict) else None
            if answers is None and isinstance(payload, dict):
                answers = payload

            if not isinstance(answers, dict):
                continue

            for domain_key, qmap in answers.items():
                if not isinstance(qmap, dict):
                    continue

                dk = str(domain_key).strip()
                if not dk:
                    continue

                aggregated.setdefault(dk, {})

                for qid, val in qmap.items():
                    qid_str = str(qid).strip()
                    if not qid_str:
                        continue
                    try:
                        score = float(val)
                    except Exception:
                        continue

                    aggregated[dk].setdefault(qid_str, [])
                    aggregated[dk][qid_str].append(score)

        # finalize
        out: Dict[str, Dict[str, float]] = {}
        for dk, qmap in aggregated.items():
            out[dk] = {}
            for qid, values in qmap.items():
                if not values:
                    continue
                # non-admin will still average (single value)
                out[dk][qid] = sum(values) / float(len(values))

        return out, included_users

    def _get_user_meta(self, user_id: str) -> UserMeta:
        users = self.repo.fetch_all("users") or []
        from auth.crypto_service import decrypt_text

        uid = str(user_id or "").strip()
        for u in users:
            if str(u.get("email_hash") or "").strip() != uid:
                continue
            try:
                full_name = decrypt_text(u.get("full_name_encrypted"))
            except Exception:
                full_name = ""
            try:
                email = decrypt_text(u.get("email_encrypted"))
            except Exception:
                email = ""
            return UserMeta(user_id=uid, full_name=full_name, email=email)

        return UserMeta(user_id=uid, full_name="", email="")

    def _get_project_name(self, project_id: str) -> str:
        projects = self.repo.fetch_all("projects") or []
        pid = str(project_id)
        for p in projects:
            if str(p.get("project_id")) == pid:
                return str(p.get("name") or "").strip()
        return ""

    # =========================================================
    # Scores
    # =========================================================

    def _compute_scores(self, domain_metas: Dict[str, DomainMeta], answers_by_domain: Dict[str, Dict[str, float]]) -> List[DomainScore]:
        """
        Only domains present in answers_by_domain are included.
        Ordering: execution_request order via domain_metas insertion order (domain_0..N created in order).
        """
        scores: List[DomainScore] = []

        for domain_key in domain_metas.keys():
            if domain_key not in answers_by_domain:
                continue

            meta = domain_metas[domain_key]
            qmap = answers_by_domain.get(domain_key, {}) or {}

            qs: List[Tuple[str, str, float]] = []
            for qid, val in qmap.items():
                qid_str = str(qid).strip()
                qtext = meta.qtext.get(qid_str.lower(), "")
                qs.append((qid_str, qtext, float(val)))

            qs.sort(key=lambda x: self._qid_sort_key(x[0]))

            if qs:
                vals = [v for _, _, v in qs]
                avg_raw = sum(vals) / float(len(vals))
                avg_floor = int(math.floor(avg_raw))
            else:
                avg_raw = 0.0
                avg_floor = 0

            scores.append(DomainScore(
                domain_key=domain_key,
                acronym=meta.acronym,
                name=meta.name,
                question_scores=qs,
                avg_raw=avg_raw,
                avg_floor=avg_floor
            ))

        return scores

    def _qid_sort_key(self, qid: str):
        s = str(qid).strip().lower()
        m = re.search(r"(\d+)", s)
        return int(m.group(1)) if m else 9999

    # =========================================================
    # Dependencies
    # =========================================================

    def _load_dependency_issues(self) -> List[DependencyIssue]:
        p = self.base_dir / "data" / "domains" / "us" / "Dependencies_inconsistencies_theory_cluster_output.json"
        data = self._safe_load_json(p)
        if not data:
            return []

        if isinstance(data, list):
            items = data
        elif isinstance(data, dict):
            items = data.get("items") or data.get("data") or []
        else:
            items = []

        out: List[DependencyIssue] = []
        for it in items:
            if not isinstance(it, dict):
                continue
            out.append(DependencyIssue(
                domain_acronym=str(it.get("domain_acronym") or "").strip(),
                reference_acronym=str(it.get("reference_acronym") or "").strip(),
                dependency_broken=bool(it.get("dependency_broken", False)),
                severity_rationale=str(it.get("severity_rationale") or "").strip(),
                scenarios=it.get("scenarios") or {}
            ))
        return out

    # =========================================================
    # Blueprint: Action + DEMO
    # =========================================================

    def _triggered_action_code(self, acronym: str, maturity_grade: int) -> str:
        return f"{acronym}-{maturity_grade + 1:02d}"

    def _load_demo_item_for_action(self, acronym: str, action_code: str, language: str):

        lang = (language or "us").lower()

        p = (
            self.base_dir
            / "data"
            / "domains"
            / lang
            / f"{acronym}_theory_demo_output_PATCHED.json"
        )

        if not p.exists():
            return None

        data = self._safe_load_json(p)
        if not data:
            return None

        items = data.get("items") if isinstance(data, dict) else data
        if not isinstance(items, list):
            return None

        for it in items:
            if str(it.get("action_code") or "").strip() == action_code:
                return it

        return None

    # =========================================================
    # DOCX helpers
    # =========================================================
    
    def _normalize_maturity_label_text(
        self,
        text: str,
        tree_data: Dict[str, Any]
    ) -> str:
        """
        Normaliza qualquer ocorrência de:
        Nível X - Algo
        Pontuação X - Algo
        para o label oficial definido no maturity_scale.
        """

        if not text or not isinstance(tree_data, dict):
            return text

        maturity_scale = tree_data.get("maturity_scale") or {}
        if not isinstance(maturity_scale, dict):
            return text

        official = {}
        for k, v in maturity_scale.items():
            try:
                level = int(k)
            except Exception:
                continue

            label = str(v.get("maturity_level") or "").strip()
            if label:
                official[level] = label

        if not official:
            return text

        import re

        def replace(match):
            level = int(match.group(1))
            if level in official:
                return f"Nível {level} - {official[level]}"
            return match.group(0)

        pattern1 = re.compile(r"Nível\s+(\d+)\s*-\s*([^\n]+)")
        text = pattern1.sub(replace, text)

        pattern2 = re.compile(r"Pontuação\s+(\d+)\s*-\s*([^\n]+)")
        text = pattern2.sub(replace, text)

        return text

    def _add_structured_kv_block(
        self,
        doc: Document,
        key: str,
        value: str
    ):
        """
        KV estruturado com espaçamento executivo.
        Usado para blocos principais como:
        Número do Procedimento
        Nome
        Pré-requisito
        Entregável
        """

        from docx.shared import Pt

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        p.add_run(f"{key}: ").bold = True
        p.add_run(value or "")
        
        
    def _render_text_with_bold_prefix_if_colon(
        self,
        doc: Document,
        text: str
    ):
        if not text:
            return

        text = text.strip()

        # Caso 1: termina com ":" → label puro
        if text.endswith(":") and ":" not in text[:-1]:
            p = doc.add_paragraph()
            p.add_run(text).bold = True
            return

        # Caso 2: formato "Label: Conteúdo"
        if ":" in text:
            prefix, suffix = text.split(":", 1)

            if len(prefix.strip()) >= 2:
                p = doc.add_paragraph()
                p.add_run(prefix.strip() + ": ").bold = True
                p.add_run(suffix.strip())
                return

        # Caso padrão
        self._add_paragraph(doc, text)
    
    def _setup_doc_styles(self, doc: Document):

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        sizes = {
            1: 20,
            2: 16,
            3: 14,
            4: 12
        }

        for lvl, size in sizes.items():
            st = doc.styles[f"Heading {lvl}"]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.bold = True

    def _add_page_break(self, doc: Document):
        doc.add_page_break()

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        doc.add_heading(text, level=level)

    def _add_paragraph(self, doc: Document, text: str):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        
    def _resolve_action_catalog_path(self, domain_acronym: str, language: str) -> Optional[Path]:
        """
        Tries to locate <DOMAIN>_action_catalog.yaml in common repo locations.
        We keep this very defensive to avoid breaking existing behavior.
        """
        dom = (domain_acronym or "").strip()
        lang = (language or "us").strip().lower()

        if not dom:
            return None

        filenames = [
            f"{dom}_action_catalog.yaml",
            f"{dom}_action_catalog.yml",
        ]

        candidates: List[Path] = []

        # Most common in your project structure (same pattern used for decision_tree)
        for fn in filenames:
            candidates.append(self.base_dir / "data" / "domains" / "Language" / lang / fn)
            candidates.append(self.base_dir / "data" / "domains" / "Language" / "us" / fn)

            # Legacy / alternate layouts
            candidates.append(self.base_dir / "data" / "domains" / lang / fn)
            candidates.append(self.base_dir / "data" / "domains" / "us" / fn)

            # Last resort (project root)
            candidates.append(self.base_dir / fn)

        for p in candidates:
            if p.exists():
                return p

        return None

    def _load_action_catalog(self, domain_acronym: str, language: str) -> Optional[Dict[str, Any]]:
        p = self._resolve_action_catalog_path(domain_acronym, language)
        if not p:
            return None
        data = self._safe_load_yaml(p)
        return data if isinstance(data, dict) else None

    def _add_procedure_elements_of_example(
        self,
        doc: Document,
        domain_acronym: str,
        action_code: str,
        language: str
    ):
        """
        Adds a new subsection "Elementos de exemplo" listing ALL procedures
        from action_catalog.yaml for the given domain+action.
        Format: titles before ':' in bold.
        """
        catalog = self._load_action_catalog(domain_acronym, language)
        if not catalog:
            return

        action = (catalog.get("action_catalog") or {}).get(action_code) or {}
        procedures = action.get("procedures") or []
        if not isinstance(procedures, list) or not procedures:
            return

        # New subsection
        self._add_heading(doc, self._t("example_elements_title", language), level=4)

        def add_kv(key: str, val: str):
            self._add_structured_kv_block(doc, key, val)

        for i, proc in enumerate(procedures, start=1):
            if not isinstance(proc, dict):
                continue
                
            is_proc6 = str(action_code).endswith("06") and int(proc.get("number") or 0) == 6
            max_bullets_proc6 = 5  # ajuste se quiser 3/4
            
            # Small separation between procedures (not a page break)
            if i > 1:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(12)

            add_kv(self._t("procedure_number", language), str(proc.get("number") or "").strip())
            add_kv(self._t("procedure_name", language), str(proc.get("name") or "").strip())
            add_kv(self._t("procedure_prerequisite", language), str(proc.get("prerequisite") or "").strip())
            add_kv(
                self._t("procedure_deliverable", language),
                str(proc.get("deliverable") or "").strip()
            )

            recs = proc.get("recommendations") or []
            if isinstance(recs, list) and recs:
                p = doc.add_paragraph()
                p.add_run(self._t("procedure_recommendations", language)).bold = True

                recs_to_render = recs[:max_bullets_proc6] if is_proc6 else recs

                for r in recs_to_render:
                    rr = str(r or "").strip()
                    if rr:
                        doc.add_paragraph(rr, style="List Bullet")

            note_value = proc.get("note") or proc.get("notes")
            if note_value:
                p = doc.add_paragraph()
                p.add_run(self._t("procedure_notes", language)).bold = True

                if isinstance(note_value, list):
                    notes_to_render = note_value[:max_bullets_proc6] if is_proc6 else note_value
                    for n in notes_to_render:
                        nn = str(n or "").strip()
                        if nn:
                            doc.add_paragraph(nn, style="List Bullet")
                else:
                    txt = str(note_value).strip()
                    if txt:
                        doc.add_paragraph(txt, style="List Bullet")
    
    
    def _add_toc_field(self, doc: Document):
        p = doc.add_paragraph()
        r = p.add_run()

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'

        fld_separate = OxmlElement('w:fldChar')
        fld_separate.set(qn('w:fldCharType'), 'separate')

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        r._r.append(fld_begin)
        r._r.append(instr)
        r._r.append(fld_separate)
        r._r.append(fld_end)

        doc.add_paragraph("")

        #doc.add_paragraph("Update the Table of Contents in Word (Right click → Update field) after opening the document.")

    def _add_cover(
        self,
        doc: Document,
        project_id: str,
        project_name: str,
        user_meta: Optional[UserMeta],
        is_admin: bool,
        included_users: List[UserMeta],
        language: str
    ):
        
        T = lambda k: self._t(k, language)
        
        title = doc.add_paragraph()
        run = title.add_run(T("cover_title_admin") if is_admin else T("cover_title_user"))
        run.bold = True
        run.font.size = Pt(28)
        

        subtitle = doc.add_paragraph()
        run2 = subtitle.add_run(T("cover_subtitle"))
        run2.font.size = Pt(16)

        doc.add_paragraph("")

        p1 = doc.add_paragraph()
        p1.add_run(f"{T('project')} ").bold = True
        p1.add_run(project_name or "")
        if project_name:
            p1.add_run(" ")
        p1.add_run(f"(ID: {project_id})")

        if is_admin:
            p2 = doc.add_paragraph()            
            p2.add_run(f"{T('scope')} ").bold = True
            p2.add_run(T("scope_all_users"))

            p3 = doc.add_paragraph()            
            p3.add_run(f"{T('users_included')} ").bold = True
            p3.add_run(str(len(included_users)))

        else:
            p2 = doc.add_paragraph()            
            p2.add_run(f"{T('assessed_user')} ").bold = True
            
            if user_meta and (user_meta.full_name or user_meta.email):
                p2.add_run(user_meta.full_name or "")
                if user_meta.email:
                    p2.add_run(f" ({user_meta.email})")
                p2.add_run(f" [ID: {user_meta.user_id}]")
            else:
                p2.add_run(f"[ID: {user_meta.user_id if user_meta else ''}]")

        p4 = doc.add_paragraph()        
        p4.add_run(f"{T('report_language')} ").bold = True
        p4.add_run(str(language))

        p5 = doc.add_paragraph()        
        p5.add_run(f"{T('generated')} ").bold = True
        p5.add_run(datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"))

    # =========================================================
    # DOCX sections
    # =========================================================

    def _add_results_section(self, doc: Document, scores: List[DomainScore], is_admin: bool, language: str):
        
        T = lambda k: self._t(k, language)        
                       
        # -------------------------------------------------
        # 1.1 Radar Chart
        # -------------------------------------------------
        
        self._add_radar_chart(doc, scores, language)                 
                
        # -------------------------------------------------
        # 1.1 Tabela por dominio 
        # -------------------------------------------------
        #doc.add_page_break()
        doc.add_paragraph("")
        self._add_heading(doc, T("domain_summary_title"), level=2)

        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = T("table_domain")
        hdr[1].text = T("table_domain_name")
        hdr[2].text = T("table_avg")
        hdr[3].text = T("table_final_grade")
        hdr[4].text = T("table_maturity_level")

        for s in scores:
            row = table.add_row().cells
            row[0].text = s.acronym
            row[1].text = s.name
            row[2].text = f"{s.avg_raw:.2f}"
            row[3].text = str(s.avg_floor)
            row[4].text = self._likert_label(s.avg_floor, language)
             
        
        for idx, s in enumerate(scores, start=2):
            doc.add_paragraph("")
            self._add_heading(doc, f"1.{idx} {s.acronym} · {s.name}", level=2)

            self._add_paragraph(
                doc,
                f"{T('domain_maturity_rule')} {s.avg_floor} ({self._likert_label(s.avg_floor,language)})."
            )

            if not s.question_scores:                
                self._add_paragraph(doc, T("no_scored_questions"))
                continue

            t = doc.add_table(rows=1, cols=4)
            h = t.rows[0].cells
            h[0].text = T("table_question")
            h[1].text = T("table_score")
            h[2].text = T("table_nearest_level")
            h[3].text = T("table_notes")

            for qid, qtext, val in s.question_scores:
                r = t.add_row().cells
                r[0].text = qtext or qid
                r[1].text = f"{val:.2f}" if is_admin else str(int(round(val)))
                nearest = int(round(val))
                if nearest < 0:
                    nearest = 0
                if nearest > 5:
                    nearest = 5
                r[2].text = self._likert_label(nearest, language)
                r[3].text = ""
            
            self._add_paragraph(doc, T("scoring_rule"))
            
            # -------------------------------------------------
            # Consistency Alert (i18n)
            # -------------------------------------------------
            vals = [v for _, _, v in s.question_scores]

            if vals:
                if max(vals) - min(vals) >= 3:

                    alert_title = T("maturity_dispersion_title")
                    alert_text = T("maturity_dispersion_text")

                    if alert_title and alert_title != "maturity_dispersion_title":
                        doc.add_paragraph("")
                        p = doc.add_paragraph()
                        run = p.add_run(f"⚠ {alert_title} ")
                        run.bold = True

                        if alert_text and alert_text != "maturity_dispersion_text":
                            p.add_run(alert_text)

    
    
    def _norm_title(self, s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").strip().lower())
        
    def _clean_placeholders(self, text: str) -> str:
        """
        Remove tokens internos tipo __ACR12__ ou padrões complexos.
        """
        if not text:
            return text

        # remove __TOKEN__
        text = re.sub(r"__[^_]+__", "", text)

        return text.strip()

    def _is_domain_context_section(self, sec_title_norm: str) -> bool:
        # "Domain and Context Framing" / "Domínio e Contexto"
        return ("domain and context" in sec_title_norm) or ("domínio e contexto" in sec_title_norm) or ("dominio e contexto" in sec_title_norm)

    def _is_action_definition_section(self, sec_title_norm: str) -> bool:
        # "Action Definition" / "Definição de Ação"
        return ("action definition" in sec_title_norm) or ("definição de ação" in sec_title_norm) or ("definicao de acao" in sec_title_norm)

    def _is_procedure_definition_section(self, sec_title_norm: str) -> bool:
        # "Procedure Definition" / "Definição de Procedimento"
        return ("procedure definition" in sec_title_norm) or ("definição de procedimento" in sec_title_norm) or ("definicao de procedimento" in sec_title_norm)

    def _is_step_by_step_subsection(self, sub_title_norm: str) -> bool:
        # "Step-by-Step Execution Structure" / "Estrutura de Execução Passo a Passo"
        return ("step-by-step" in sub_title_norm) or ("passo a passo" in sub_title_norm) or ("execução passo" in sub_title_norm) or ("execucao passo" in sub_title_norm)

    def _load_domain_tree_and_catalog(
        self,
        project_id: str,
        domain_acronym: str,
        flow: Dict[str, Any],
        language: str
    ) -> Tuple[Optional[Dict[str, Any]], Optional[Dict[str, Any]]]:
        """
        Resolve paths no padrão do projeto (data/projects/<id>/Domains/<lang>/...)
        Usando flow.yaml Domain_flow -> files: decision_tree / action_catalog.
        """
        lang = (language or "us").strip().lower() or "us"

        domain_flow = flow.get("Domain_flow", []) or []
        dom = next((d for d in domain_flow if str(d.get("acronym") or "").strip().upper() == str(domain_acronym).strip().upper()), None) or {}
        files = dom.get("files") or {}

        decision_tree_rel = files.get("decision_tree")
        action_catalog_rel = files.get("action_catalog")

        base = self.base_dir / "data" / "projects" / str(project_id) / "Domains"

        candidates_lang = [lang, "us"]
        tree_data = None
        catalog_data = None

        for lg in candidates_lang:
            if decision_tree_rel and tree_data is None:
                p = base / lg / str(decision_tree_rel)
                if p.exists():
                    tree_data = self._safe_load_yaml(p) or None

            if action_catalog_rel and catalog_data is None:
                p = base / lg / str(action_catalog_rel)
                if p.exists():
                    catalog_data = self._safe_load_yaml(p) or None

            if tree_data is not None and catalog_data is not None:
                break

        return tree_data, catalog_data

    def _first_question_node(self, tree_data: Optional[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        if not isinstance(tree_data, dict):
            return None
        questions = tree_data.get("questions") or {}
        if not isinstance(questions, dict) or not questions:
            return None

        # pega a primeira por ordem Q1, Q2... se possível
        def qkey(k: str) -> int:
            m = re.search(r"(\d+)", str(k))
            return int(m.group(1)) if m else 9999

        first_key = sorted(questions.keys(), key=qkey)[0]
        qnode = questions.get(first_key)
        return qnode if isinstance(qnode, dict) else None

    def _domain_description(self, tree_data: Optional[Dict[str, Any]]) -> str:
        if not isinstance(tree_data, dict):
            return ""
        dom = tree_data.get("domain") or {}
        if isinstance(dom, dict):
            return str(dom.get("description") or "").strip()
        return ""

    def _domain_objectives(self, tree_data: Optional[Dict[str, Any]]) -> List[str]:
        if not isinstance(tree_data, dict):
            return []
        questions = tree_data.get("questions") or {}
        if not isinstance(questions, dict):
            return []
        out = []
        for _, qinfo in questions.items():
            if not isinstance(qinfo, dict):
                continue
            obj = str(qinfo.get("objective") or "").strip()
            if obj:
                out.append(obj)
        return out

    def _action_from_tree_for_score(self, tree_data: Optional[Dict[str, Any]], score_floor: int) -> Tuple[str, str]:
        """
        Retorna (action_code, description) a partir do score_action_mapping do decision_tree.
        Assume consistência do mapping entre questions.
        """
        qnode = self._first_question_node(tree_data)
        if not isinstance(qnode, dict):
            return "", ""

        mapping = qnode.get("score_action_mapping") or {}
        if not isinstance(mapping, dict):
            return "", ""

        action_node = mapping.get(score_floor)
        if not isinstance(action_node, dict):
            # às vezes o YAML serializa keys como string
            action_node = mapping.get(str(score_floor))

        if not isinstance(action_node, dict):
            return "", ""

        ac = str(action_node.get("action_code") or "").strip()
        desc = str(action_node.get("description") or "").strip()
        return ac, desc

    def _find_action_in_catalog(self, catalog_data: Optional[Dict[str, Any]], action_code: str) -> Optional[Dict[str, Any]]:
        if not isinstance(catalog_data, dict):
            return None
        actions = catalog_data.get("actions")
        if isinstance(actions, dict):
            node = actions.get(action_code)
            return node if isinstance(node, dict) else None
        if isinstance(actions, list):
            for it in actions:
                if isinstance(it, dict) and str(it.get("action_code") or "").strip() == action_code:
                    return it
        return None

    def _add_maturity_levels_section(
        self,
        doc: Document,
        tree_data: Dict[str, Any],
        language: str
    ):
        """
        Renderiza maturity_scale completo do decision_tree.yaml
        """
        
        T = lambda k: self._t(k, language)
        
        maturity_scale = tree_data.get("maturity_scale")
        if not isinstance(maturity_scale, dict):
            return

        def sort_key(k):
            try:
                return int(k)
            except Exception:
                return 999

        for level_key in sorted(maturity_scale.keys(), key=sort_key):

            level_data = maturity_scale.get(level_key)
            if not isinstance(level_data, dict):
                continue

            maturity_level = level_data.get("maturity_level", "")
            meaning = level_data.get("meaning", "")
            interpretation = level_data.get("interpretation_for_assessment", "")

            # Heading por nível
            self._add_heading(
                doc,
                f"{T('maturity_level_prefix')} {level_key} – {maturity_level}",
                level=3
            )

            if meaning:
                self._add_paragraph(doc, meaning)

            if interpretation:
                self._add_paragraph(doc, interpretation)

    def _inject_domain_and_context(self, doc: Document, tree_data: Optional[Dict[str, Any]]):
        """
        3.x.3 Definição de Domínio e Contexto:
          - 1 parágrafo: decision_tree.domain.description
          - bullets: todos os objectives das questions
        """
        desc = self._domain_description(tree_data)
        if desc:
            self._add_paragraph(doc, desc)

        objectives = self._domain_objectives(tree_data)
        if objectives:
            for obj in objectives:
                doc.add_paragraph(obj, style="List Bullet")

    def _inject_action_definition(self, doc: Document, tree_data: Optional[Dict[str, Any]], score_floor: int):
        """
        3.x.4 Definição de Ação:
          - 1 parágrafo: action_code + description vindo do score_action_mapping do decision_tree
        """
        ac, desc = self._action_from_tree_for_score(tree_data, score_floor)
        if not ac and not desc:
            return

        if desc:
            self._add_paragraph(doc, f"{ac}: {desc}" if ac else desc)
        else:
            self._add_paragraph(doc, f"{ac}")
    
                
    def _add_maturity_levels_block(
        self,
        doc: Document,
        project_id: str,
        scores: List[DomainScore],
        flow: Dict[str, Any],
        language: str
    ):
        """
        Renderiza 3.1 Maturity Levels com base no maturity_scale
        do decision_tree.yaml do primeiro domínio.
        Não altera nada do fluxo existente.
        """

        if not scores:
            return

        domain_acronym = scores[0].acronym

        tree_data, _ = self._load_domain_tree_and_catalog(
            project_id=project_id,
            domain_acronym=domain_acronym,
            flow=flow,
            language=language
        )

        if not isinstance(tree_data, dict):
            return

        maturity_scale = tree_data.get("maturity_scale")
        if not isinstance(maturity_scale, dict):
            return

        # Heading principal
        self._add_heading(doc, T("maturity_levels_title"), level=2)

        # Ordenar níveis numericamente
        def sort_key(k):
            try:
                return int(k)
            except Exception:
                return 999

        for level_key in sorted(maturity_scale.keys(), key=sort_key):

            level_data = maturity_scale.get(level_key)
            if not isinstance(level_data, dict):
                continue

            maturity_level = level_data.get("maturity_level", "")
            meaning = level_data.get("meaning", "")
            interpretation = level_data.get("interpretation_for_assessment", "")

            # Subheading por nível
            self._add_heading(
                doc,
                f"{T('maturity_level_prefix')} {level_key} – {maturity_level}",
                level=3
            )

            if meaning:
                self._add_paragraph(doc, meaning)

            if interpretation:
                self._add_paragraph(doc, interpretation)
                
    
    def _add_blueprint_section(
        self,
        doc: Document,
        project_id: str,
        scores: List[DomainScore],
        domain_metas: Dict[str, DomainMeta],
        flow: Dict[str, Any],
        language: str
    ):
        T = lambda k: self._t(k, language)
        
        tree_data = None
        
        

        # --- LOAD DEMO CODE MAP ---
        demo_map_path = self.base_dir / "data" / "general" / "demo_code_map.yaml"
        demo_map = self._safe_load_yaml(demo_map_path) or {}
        struct_codes = demo_map.get("demo_code_map", {}).get("structural_codes", {})

        CODE_DOMAIN_SECTION = str(struct_codes.get("domain_context_section", "3"))
        CODE_ORG_CONTEXT = str(struct_codes.get("organizational_context_subsection", "3.2"))
        CODE_EXECUTION_STRUCTURE = str(struct_codes.get("execution_structure_subsection", "5.4"))

        lang = (language or "us").lower()
        texts = self._report_texts or {}

        global_keywords = texts.get(lang, {}).get("blueprint_global_keywords", [])
        common_template_keywords = texts.get(lang, {}).get("blueprint_common_template_keywords", [])
        closure_keywords = texts.get(lang, {}).get("blueprint_closure_keywords", [])

        def is_global_section(title: str) -> bool:
            t = self._norm_title(title)
            return any(self._norm_title(k) in t for k in global_keywords)

        def is_common_template(title: str) -> bool:
            t = self._norm_title(title)
            return any(self._norm_title(k) in t for k in common_template_keywords)

        def is_closure(title: str) -> bool:
            t = self._norm_title(title)
            return any(self._norm_title(k) in t for k in closure_keywords)
    
        def render_demo_section_blocks(
            sec: Dict[str, Any],
            tree_data: Optional[Dict[str, Any]]
        ):

            for block in sec.get("blocks", []):

                # -----------------------------
                # Normaliza texto do bloco (se houver)
                # -----------------------------
                if "text" in block and isinstance(block["text"], str):
                    block["text"] = self._normalize_maturity_label_text(
                        block["text"],
                        tree_data
                    )

                # -----------------------------
                # SUBSECTION
                # -----------------------------
                if block["kind"] == "subsection":

                    self._add_heading(doc, block["title"], level=4)

                    for item in block.get("items", []):

                        # Normaliza texto do item
                        if "text" in item and isinstance(item["text"], str):
                            item["text"] = self._normalize_maturity_label_text(
                                item["text"],
                                tree_data
                            )

                        if item["type"] == "bullet":
                            doc.add_paragraph(item["text"], style="List Bullet")

                        elif item["type"] == "kv":
                            p = doc.add_paragraph()
                            p.add_run(item["key"] + ": ").bold = True
                            p.add_run(item["text"])

                        else:
                            self._render_text_with_bold_prefix_if_colon(
                                doc,
                                item.get("text", "")
                            )

                # -----------------------------
                # KV
                # -----------------------------
                elif block["kind"] == "kv":

                    p = doc.add_paragraph()
                    p.add_run(block["key"] + ": ").bold = True
                    p.add_run(block.get("text", ""))

                # -----------------------------
                # BULLET
                # -----------------------------
                elif block["kind"] == "bullet":

                    doc.add_paragraph(block.get("text", ""), style="List Bullet")

                # -----------------------------
                # TEXTO PADRÃO
                # -----------------------------
                else:

                    self._render_text_with_bold_prefix_if_colon(
                        doc,
                        block.get("text", "")
                    )

        global_sections0: List[Dict[str, Any]] = []
        sections0_all: List[Dict[str, Any]] = []

        if scores:
            s0 = scores[0]
            action0 = self._triggered_action_code(s0.acronym, s0.avg_floor)
            item0 = self._load_demo_item_for_action(s0.acronym, action0, language)
            demo0 = str(item0.get("demo") or "").strip() if item0 else ""
            sections0 = self._parse_demo_sections(demo0) if demo0 else []
            sections0_all = sections0
            global_sections0 = [sec for sec in sections0_all if is_global_section(sec.get("title", ""))]

        # -------------------------------------------------
        # 3.1 Assessment Result and Strategic Direction
        # -------------------------------------------------
        self._add_heading(doc, T("section_3_assessment_result"), level=2)
        self._add_paragraph(doc, T("blueprint_context"))

        if scores:
            self._add_paragraph(doc, T("executive_reading"))
            for s in scores:
                doc.add_paragraph(
                    f"{s.acronym} = {s.avg_raw:.2f} → "
                    f"{T('maturity_level_prefix')} {s.avg_floor} "
                    f"({self._likert_label(s.avg_floor, language)})",
                    style="List Bullet"
                )

        # -------------------------------------------------
        # 3.2 Common Elements
        # -------------------------------------------------
        self._add_heading(doc, T("section_3_common_elements"), level=2)
        self._add_paragraph(doc, T("section_3_common_elements_intro"))
    
        # 3.2.1
        self._add_heading(doc, T("section_3_common_template"), level=3)
        self._add_paragraph(doc, T("section_3_common_template_intro"))

        sec_counter = 1
        for sec in global_sections0:
            title_norm = self._norm_title(sec.get("title", ""))
            if not is_common_template(sec.get("title", "")):
                continue

            section_number = f"{T('section_3_common_template').split()[0]}.{sec_counter}"
            sec_counter += 1

            self._add_heading(doc, f"{section_number} {sec['title']}", level=4)
            render_demo_section_blocks(sec, tree_data)

        # 3.2.2
        self._add_heading(doc, T("section_3_common_transversal"), level=3)
        self._add_paragraph(doc, T("section_3_common_transversal_intro"))

        sec_counter = 1
        for sec in sections0_all:
            title_norm = self._norm_title(sec.get("title", ""))
            if is_common_template(sec.get("title", "")):
                continue
            if is_closure(sec.get("title", "")):
                continue

            section_number = f"{T('section_3_common_transversal').split()[0]}.{sec_counter}"
            sec_counter += 1

            self._add_heading(doc, f"{section_number} {sec['title']}", level=4)
            render_demo_section_blocks(sec, tree_data)

        # -------------------------------------------------
        # 3.3 Consolidated Maturity Model
        # -------------------------------------------------
        maturity_tree = None
        if scores:
            maturity_tree, _ = self._load_domain_tree_and_catalog(
                project_id=project_id,
                domain_acronym=scores[0].acronym,
                flow=flow,
                language=language
            )

        self._add_heading(doc, T("section_3_maturity_consolidated"), level=2)

        if isinstance(maturity_tree, dict) and maturity_tree.get("maturity_scale"):
            self._add_maturity_levels_section(doc, maturity_tree, language)
        else:
            txt = T("no_maturity_scale")
            if txt and txt != "no_maturity_scale":
                self._add_paragraph(doc, txt)

        # -------------------------------------------------
        # Domain sections
        # -------------------------------------------------
        for idx, s in enumerate(scores, start=1):

            chapter_number = f"3.{idx + 3}"

            # 3.x Domain Heading
            self._add_heading(
                doc,
                f"{chapter_number} {s.acronym} · {s.name}",
                level=2
            )
            
            # -------------------------------
            # Load decision tree
            # -------------------------------
            tree_data, _ = self._load_domain_tree_and_catalog(
                project_id=project_id,
                domain_acronym=s.acronym,
                flow=flow,
                language=language
            )


            # -------------------------------
            # Triggered Action (primeiro define tudo)
            # -------------------------------
            action_code = self._triggered_action_code(s.acronym, s.avg_floor)

            item = self._load_demo_item_for_action(
                s.acronym,
                action_code,
                language
            )

            title = ""
            demo = ""

            if item:
                title = str(
                    item.get("action_title")
                    or item.get("title")
                    or item.get("name")
                    or ""
                ).strip()
                demo = str(item.get("demo") or "").strip()

            
            demo = self._normalize_maturity_labels_in_demo(demo, tree_data)
            
            
            # -------------------------------
            # 3.x.1 Definição de Domínio e Contexto
            # -------------------------------
            context_number = f"{chapter_number}.1"

            # evitar duplicação lógica com 3.2.2.1
            lbl = T("section_3_domain_context_domain_specific_title")

            if (
                not lbl
                or lbl == "section_3_domain_context_domain_specific_title"
                or str(lbl).startswith("[section_")
            ):
                lbl = "Domain Contextualization (Assessed Domain)"

            self._add_heading(
                doc,
                f"{context_number} {lbl}",
                level=3
            )

            # descrição do domínio
            desc = self._domain_description(tree_data)
            if desc:
                p = doc.add_paragraph()
                p.add_run(T("description_label")).bold = True
                p.add_run(desc)
                        
            self._render_text_with_bold_prefix_if_colon(
                doc,
                T("section_3_domain_context_objective")
            )

            self._render_text_with_bold_prefix_if_colon(
                doc,
                T("section_3_domain_context_expected")
            )

            self._add_paragraph(
                doc,
                f"{T('current_maturity')} "
                f"{T('maturity_level_prefix')} {s.avg_floor} "
                f"({self._likert_label(s.avg_floor, language)})."
            )

            # -------------------------------
            # Contexto Organizacional (usa demo já definido)
            # -------------------------------
            if demo:

                parsed = self._parse_demo_sections(demo)

                sec3 = next(
                    (x for x in parsed
                     if str(x.get("code") or "").strip() == CODE_DOMAIN_SECTION),
                    None
                )

                if sec3:

                    sub32 = None
                    for b in sec3.get("blocks", []):
                        if (
                            b.get("kind") == "subsection"
                            and str(b.get("code") or "").strip() == CODE_ORG_CONTEXT
                        ):
                            sub32 = b
                            break

                    if sub32:
                        self._add_heading(doc, sub32["title"], level=4)

                        for it in sub32.get("items", []):
                            if it["type"] == "bullet":
                                doc.add_paragraph(it["text"], style="List Bullet")

                            elif it["type"] == "kv":
                                p = doc.add_paragraph()
                                p.add_run(it["key"] + ": ").bold = True
                                p.add_run(it["text"])

                            else:
                                self._render_text_with_bold_prefix_if_colon(
                                    doc,
                                    it["text"]
                                )

            # -------------------------------
            # Triggered Action Label
            # -------------------------------
            if title:
                self._add_paragraph(
                    doc,
                    f"{T('triggered_action')} {action_code} · {title}"
                )
            else:
                self._add_paragraph(
                    doc,
                    f"{T('triggered_action_code')} {action_code}"
                )

            if not demo:
                self._add_paragraph(doc, T("no_demo"))
                continue

            # -------------------------------
            # 3.x.2 Pacote de Procedimentos
            # -------------------------------
            procedure_number = f"{chapter_number}.2"

            self._add_heading(
                doc,
                f"{procedure_number} {T('procedure_pack_full')}",
                level=3
            )

            self._add_paragraph(doc, T("section_3_procedure_pack_intro"))

            sections = self._parse_demo_sections(demo)

            if not sections:
                self._add_paragraph(doc, demo)
                continue

            # Remove global document-level sections
            sections = [
                sec for sec in sections
                if not is_global_section(sec.get("title", ""))
            ]
            
            # Avoid duplicating "Domain Definition and Context" (already rendered in 3.x.1)
            sections = [
                sec for sec in sections
                if str(sec.get("code") or "").strip() != CODE_DOMAIN_SECTION
            ]

            render_example_block = True

            # Render domain-specific sections
            for sec_index, sec in enumerate(sections, start=1):

                section_number = f"{procedure_number}.{sec_index}"

                title = sec["title"]
                
                is_procedure_6 = str(action_code).endswith("06")

                normalized = str(title).strip().lower()

                if normalized == T("procedure_action_definition").lower():
                    title = f"Ação {action_code}"

                elif normalized == T("procedure_definition").lower():
                    title = f"Procedimentos da Ação {action_code}"

                self._add_heading(
                    doc,
                    f"{section_number} {title}",
                    level=4
                )

                for block in sec["blocks"]:

                    if block["kind"] == "subsection":
                        
                        if str(block["title"]).strip().lower() == "elementos de exemplo":
                            self._add_heading(doc, block["title"], level=6)
                            continue

                        self._add_heading(doc, block["title"], level=5)

                        # Injeta os elementos prescritivos ANTES do conteúdo normal (topo)
                        if (
                            render_example_block
                            and str(block.get("code") or "").strip() == CODE_EXECUTION_STRUCTURE
                        ):
                            self._add_procedure_elements_of_example(
                                doc=doc,
                                domain_acronym=s.acronym,
                                action_code=action_code,
                                language=language
                            )
                            render_example_block = False

                        # Renderiza conteúdo normal
                        for item in block["items"]:                            

                            if item["type"] == "bullet":
                                doc.add_paragraph(item["text"], style="List Bullet")

                            elif item["type"] == "kv":
                                p = doc.add_paragraph()
                                p.add_run(item["key"] + ": ").bold = True
                                p.add_run(item["text"])

                            else:
                                # melhorar títulos curtos isolados
                                if ":" not in item["text"] and len(item["text"].split()) <= 4:
                                    p = doc.add_paragraph()
                                    p.add_run(item["text"]).bold = True
                                else:
                                    self._render_text_with_bold_prefix_if_colon(
                                        doc,
                                        item["text"]
                                    )

                        # Depois injeta os elementos prescritivos
                        if (
                            T("procedure_execution_structure") in block["title"]
                            and render_example_block
                        ):
                            self._add_procedure_elements_of_example(
                                doc=doc,
                                domain_acronym=s.acronym,
                                action_code=action_code,
                                language=language
                            )
                            render_example_block = False

                    elif block["kind"] == "kv":

                        self._override_domain_definition_kv(
                            block=block,
                            domain_acronym=s.acronym,
                            domain_name=s.name
                        )

                        p = doc.add_paragraph()
                        p.add_run(block["key"] + ": ").bold = True
                        p.add_run(block["text"])

                    elif block["kind"] == "bullet":
                        doc.add_paragraph(block["text"], style="List Bullet")

                    else:
                        self._render_text_with_bold_prefix_if_colon(
                            doc,
                            block["text"]
                        )

        # -------------------------------------------------
        # Document Closure
        # -------------------------------------------------
        closing_chapter = f"3.{len(scores) + 4}"

        self._add_heading(
            doc,
            f"{closing_chapter} {T('section_3_domain_specific_closure')}",
            level=2
        )

        self._add_heading(
            doc,
            f"{closing_chapter}.1 {T('section_3_document_governance')}",
            level=3
        )
        self._add_paragraph(doc, T("section_3_document_governance_intro"))

        self._add_heading(
            doc,
            f"{closing_chapter}.2 {T('section_3_minimum_evidence')}",
            level=3
        )
        self._add_paragraph(doc, T("section_3_minimum_evidence_intro"))

        self._add_heading(
            doc,
            f"{closing_chapter}.3 {T('section_3_evolution_criteria')}",
            level=3
        )
        self._add_paragraph(doc, T("section_3_evolution_criteria_intro"))

    def _parse_demo_sections(self, demo: str) -> List[Dict[str, Any]]:
        text = (demo or "").replace("\r\n", "\n").strip()
        if not text:
            return []

        # Detecta títulos no formato "1. Title"
        pat = re.compile(r"(?m)^\s*(\d+)\.\s+(.+?)\s*$")
        matches = list(pat.finditer(text))
        if not matches:
            return []

        chunks = []
        for i, m in enumerate(matches):
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            chunks.append(text[start:end].strip())

        out = []
        for ch in chunks:
            lines = ch.split("\n")
            raw_title = lines[0].strip()

            m0 = re.match(r"^\s*(\d+)\.\s+(.+?)\s*$", raw_title)
            code = m0.group(1).strip() if m0 else ""

            # REMOVE prefixo numérico tipo "1. "
            title = re.sub(r"^\s*\d+\.\s*", "", raw_title).strip()

            body = "\n".join(lines[1:]).strip()
            blocks = self._parse_demo_body_blocks(body)

            out.append({
                "code": code,
                "title": title,
                "blocks": blocks
            })

        return out

    def _parse_demo_body_blocks(self, body: str) -> List[Dict[str, Any]]:
        if not body:
            return []

        blocks = []

        for raw in body.split("\n"):
            if not raw.strip():
                continue

            raw_strip = raw.strip()

            # Captura "7.1 Title" / "5.4 Title:" etc
            mcode = re.match(r"^\s*(\d+(?:\.\d+)*)\s+(.+?)\s*$", raw_strip)
            raw_code = mcode.group(1).strip() if mcode else ""
            line = mcode.group(2).strip() if mcode else raw_strip

            line = self._clean_placeholders(line)

            # Subsection header:
            # 1) termina com ":" e só tem esse ":" no fim
            # 2) OU tem código x.y (ex: 7.1 / 7.3 / 5.4) mesmo sem ":"
            is_subsection_by_colon = line.endswith(":") and (":" not in line[:-1])
            is_subsection_by_code = bool(raw_code) and bool(re.match(r"^\d+\.\d+$", raw_code))

            if is_subsection_by_colon or is_subsection_by_code:
                title = line[:-1].strip() if is_subsection_by_colon else line.strip()
                blocks.append({
                    "kind": "subsection",
                    "code": raw_code,
                    "title": title,
                    "items": []
                })
                continue

            # Bullet explícito
            if line.startswith("- "):
                text = line[2:].strip()
                if blocks and blocks[-1]["kind"] == "subsection":
                    blocks[-1]["items"].append({"type": "bullet", "text": text})
                else:
                    blocks.append({"kind": "bullet", "text": text})
                continue

            # KV
            m = re.match(r"^([A-Za-z][A-Za-z\s]+):\s*(.+)$", line)
            if m:
                key = m.group(1).strip()
                val = m.group(2).strip()
                if blocks and blocks[-1]["kind"] == "subsection":
                    blocks[-1]["items"].append({"type": "kv", "key": key, "text": val})
                else:
                    blocks.append({"kind": "kv", "key": key, "text": val})
                continue

            # Texto simples
            if blocks and blocks[-1]["kind"] == "subsection":
                blocks[-1]["items"].append({"type": "text", "text": line})
            else:
                blocks.append({"kind": "text", "text": line})

        return blocks

    def _add_references_section(self, doc: Document, language: str):
        T = lambda k: self._t(k, language)
        refs = [
            "DAMA International. (n.d.). DAMA-DMBOK: Data Management Body of Knowledge (2nd ed.).",
            "EDM Council. (n.d.). DCAM: Data Management Capability Assessment Model.",
            "CMMI Institute. (n.d.). CMMI Model.",
            "European Union. (2016). General Data Protection Regulation (GDPR) (EU) 2016/679.",
            "DOMMx publication (Zenodo record: 18020434).",
            "SLR DOMMx publication (IEEE Xplore document: 11223201).",
        ]
        for r in refs:
            doc.add_paragraph(r, style="List Bullet")

    # =========================================================
    # PDF export
    # =========================================================

    def _export_pdf(
        self,
        pdf_path: str,
        project_id: str,
        project_name: str,
        user_meta: Optional[UserMeta],
        is_admin: bool,
        included_users: List[UserMeta],
        language: str,
        scores: List[DomainScore],
        domain_metas: Dict[str, DomainMeta],
        deps_issues: List[DependencyIssue],
    ):
        # PDF temporarily disabled
        return

    def _pdf_draw_wrapped(self, c, text: str, x: float, y: float, max_w: float, leading: float) -> float:
        words = (text or "").split()
        if not words:
            return y

        line = ""
        for w in words:
            trial = (line + " " + w).strip()
            if c.stringWidth(trial, c._fontname, c._fontsize) <= max_w:
                line = trial
            else:
                c.drawString(x, y, line)
                y -= leading
                line = w

        if line:
            c.drawString(x, y, line)
            y -= leading

        return y